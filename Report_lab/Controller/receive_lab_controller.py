from PySide6.QtWidgets import QMessageBox, QAbstractItemView
from PySide6.QtCore import QObject, Qt
from PySide6.QtGui import QStandardItemModel, QStandardItem
from View.view_receive_lab_frame import ReceiveLabFormView
from SERVICES_REPORT_LAB.receive_lab_service import ReceiveLabService



class ReceiveLabController(QObject):
    """ Controller for the Receive Lab Page """

    # ==========================================
    # INITIALIZATION - การเริ่มต้น
    # ==========================================
    def __init__(self, view: ReceiveLabFormView, main_controller=None):
        super().__init__()
        self.view: ReceiveLabFormView = view
        self.receive_lab_service = ReceiveLabService()
        self.main_controller = main_controller
        # Lazy loading state
        self.current_offset = 0
        self.limit = 50
        self.is_loading = False
        self.has_more_data = True
        self.all_data = []
        # Room and access control state
        self.log_room = None
        self.log_room_id = None
        self.admin_comein = False
        # Template selection state
        self.selected_template = None
        self.all_templates = []  # เก็บรายการ template ทั้งหมด
        self.current_test_items = []  # เก็บรายการตรวจปัจจุบัน
        # Setup table model
        self.setup_table_model()
        
        self._setup_connections()
    
    # ==========================================
    # TABLE SETUP - ตั้งค่าตาราง
    # ==========================================
    def setup_table_model(self):
        """Setup QStandardItemModel for tableView"""
        self.model = QStandardItemModel()
        self.model.setHorizontalHeaderLabels(['เวลา', 'Lab Order ID', 'ความเร็ว', 'ตัวอย่าง'])
        self.view.ui.tableView.setModel(self.model)
        self.view.ui.tableView.setColumnWidth(0, 180)   # เวลา
        self.view.ui.tableView.setColumnWidth(1, 220)   # Lab Order ID
        self.view.ui.tableView.setColumnWidth(2, 200)   # ความเร็ว
        self.view.ui.tableView.setColumnWidth(3, 240)   # ตัวอย่าง (remaining space)
        self.view.ui.tableView.horizontalHeader().setStretchLastSection(True)
        self.view.ui.tableView.setShowGrid(True)
        self.view.ui.tableView.setSelectionBehavior(self.view.ui.tableView.SelectionBehavior.SelectRows)
        self.view.ui.tableView.setSelectionMode(self.view.ui.tableView.SelectionMode.SingleSelection)
        self.view.ui.tableView.setAlternatingRowColors(True)
        
        # Setup QStandardItemModel for tableView_2 (Detail view)
        self.detail_model = QStandardItemModel()
        self.detail_model.setHorizontalHeaderLabels(['ชื่อการตรวจ', 'จำนวน'])
        self.view.ui.tableView_2.setModel(self.detail_model)
        self.view.ui.tableView_2.setColumnWidth(0, 390)  # ชื่อการตรวจ - ปรับให้กว้างขึ้น
        self.view.ui.tableView_2.setColumnWidth(1, 140)  # จำนวน - เพิ่มความกว้าง
        self.view.ui.tableView_2.horizontalHeader().setStretchLastSection(False)  # ปิด stretch เพื่อไม่ให้มี scrollbar
        self.view.ui.tableView_2.horizontalHeader().setSectionResizeMode(0, self.view.ui.tableView_2.horizontalHeader().ResizeMode.Stretch)  # ให้คอลัมน์แรก stretch
        self.view.ui.tableView_2.setShowGrid(True)
        self.view.ui.tableView_2.setSelectionBehavior(self.view.ui.tableView_2.SelectionBehavior.SelectRows)
        self.view.ui.tableView_2.setSelectionMode(self.view.ui.tableView_2.SelectionMode.SingleSelection)
        self.view.ui.tableView_2.setAlternatingRowColors(True)
        
        # Setup QStandardItemModel for tableView_3 (Template view)
        self.template_model = QStandardItemModel()
        self.template_model.setHorizontalHeaderLabels(['ชื่อไฟล์ Template'])
        self.view.ui.tableView_3.setModel(self.template_model)
        self.view.ui.tableView_3.horizontalHeader().setStretchLastSection(True)
        self.view.ui.tableView_3.setShowGrid(True)
        self.view.ui.tableView_3.setSelectionBehavior(self.view.ui.tableView_3.SelectionBehavior.SelectRows)
        self.view.ui.tableView_3.setSelectionMode(self.view.ui.tableView_3.SelectionMode.SingleSelection)
        self.view.ui.tableView_3.setAlternatingRowColors(True)
    
    # ==========================================
    # SIGNAL CONNECTIONS - การเชื่อมต่อสัญญาณ
    # ==========================================
    def _setup_connections(self):
        self.view.ui.clear_pushButton.clicked.connect(self.clear_pushButton_clicked)
        self.view.ui.export_pushButton.clicked.connect(self.export_pushButton_clicked)
        self.view.ui.search_pushButton.clicked.connect(self.loaded_lab_orders)
        self.view.ui.receive_pushButton.clicked.connect(self.receive_lab_orders)
        self.view.ui.reject_pushButton.clicked.connect(self.reject_lab_orders)
        scrollbar = self.view.ui.tableView.verticalScrollBar()
        scrollbar.valueChanged.connect(self.on_scroll)
        self.view.ui.tableView.doubleClicked.connect(self.on_cell_double_clicked)
        self.view.ui.tableView.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.view.ui.tableView.setSelectionBehavior(QAbstractItemView.SelectRows)
        # เชื่อมต่อการคลิกเลือก template
        self.view.ui.tableView_3.clicked.connect(self.on_template_selected)
        # เชื่อมต่อการคลิกเลือก template
        self.view.ui.tableView_3.clicked.connect(self.on_template_selected)

    # ==========================================
    # EVENT HANDLERS - TABLE INTERACTIONS - จัดการเหตุการณ์ตาราง
    # ==========================================
    def on_cell_double_clicked(self, index):
        """เมื่อ double click ที่แถวใน tableView จะแสดงรายละเอียดของ Lab Order นั้น"""
        row = index.row()
        lab_order_id = self.model.item(row, 1).text()
        lab_order_id = lab_order_id.lstrip('0')
        
        # เก็บ lab_order_id สำหรับใช้ในการรับแลป
        self.current_lab_order_id = int(lab_order_id)
        
        if self.admin_comein == True:
            room_id = self.view.get_type_search()
        else:
            room_id = self.log_room_id if hasattr(self, 'log_room_id') else None
            
        print(f"Lab Order ID: {lab_order_id} from Room ID: {room_id}")
        
        # เรียก API เพื่อดึงรายละเอียด Lab Order
        try:
            result = self.receive_lab_service.get_lab_order_details(lab_order_id, str(room_id))
            
            if result and result.get('success', False):
                order_data = result.get('order_data', {})
                test_items = result.get('test_items', [])
                
                # แสดงรายละเอียด Lab Order
                self.display_lab_order_details(order_data, test_items)
            else:
                message = result.get('message', 'ไม่สามารถดึงข้อมูลได้') if result else 'ไม่สามารถดึงข้อมูลได้'
                QMessageBox.warning(self.view, "Error", message)
                
        except Exception as e:
            QMessageBox.warning(self.view, "Error", f"ไม่สามารถดึงรายละเอียดได้: {str(e)}")
    
    def on_template_selected(self, index):
        """เมื่อคลิกเลือก template ใน tableView_3"""
        row = index.row()
        template_name = self.template_model.item(row, 0).text()
        
        # หา template ที่ตรงกันจากรายการ all_templates
        self.selected_template = None
        for template in self.all_templates:
            if template.get('report_name') == template_name:
                self.selected_template = template
                # print(f"DEBUG: เลือก template: {template_name}")
                # print(f"DEBUG: Path: {template.get('report_path')}")
                break
    
    def display_lab_order_details(self, order_data, test_items):
        """แสดงรายละเอียด Lab Order และรายการตรวจในตาราง"""
        # เคลียร์ตารางรายละเอียดเดิม
        self.detail_model.removeRows(0, self.detail_model.rowCount())
        # เคลียร์ตาราง template
        self.template_model.removeRows(0, self.template_model.rowCount())
        
        # เพิ่มรายละเอียด Lab Order (แถวแรกๆ)
        details = [
            ("=== รายละเอียด Lab Order ===", ""),
            ("Lab Order ID:", str(order_data.get('lab_order_id', '')).zfill(12)),
            ("วันที่-เวลา:", str(order_data.get('dtime', ''))),
            ("ตัวอย่างที่ส่งตรวจ:", str(order_data.get('sample_inspection', order_data.get('sample_type', '')))),
            ("ความเร็ว:", str(order_data.get('speed', ''))),
            ("", ""),
            ("=== รายการตรวจ ===", "")
        ]
        
        for detail_name, detail_value in details:
            name_item = QStandardItem(detail_name)
            value_item = QStandardItem(detail_value)
            if "===" in detail_name:
                # ทำให้แถวหัวข้อหนาขึ้น
                font = name_item.font()
                font.setBold(True)
                name_item.setFont(font)
            self.detail_model.appendRow([name_item, value_item])
        
        # เพิ่มรายการตรวจ
        if test_items:
            for idx, item in enumerate(test_items, 1):
                test_name = item.get('test_name', '')
                test_amount = item.get('test_amount', '')
                
                # แสดงชื่อการตรวจ
                name_item = QStandardItem(test_name)
                
                # แสดงจำนวน
                amount_item = QStandardItem(str(test_amount))
                amount_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                
                self.detail_model.appendRow([name_item, amount_item])
        else:
            # ไม่มีรายการตรวจ
            no_data_item = QStandardItem("ไม่พบรายการตรวจ")
            empty_item = QStandardItem("")
            self.detail_model.appendRow([no_data_item, empty_item])
        
        # เก็บ test_items ไว้สำหรับใช้หลังจากรับแลปสำเร็จ
        self.current_test_items = test_items
        # หมายเหตุ: จะแสดง template หลังจากกดปุ่มรับแลปสำเร็จเท่านั้น
    
    def find_and_display_matching_templates(self, test_items):
        """ค้นหาและแสดงไฟล์ template ที่ตรงกับรายการตรวจ"""
        if not test_items:
            print("DEBUG: ไม่มี test_items")
            return
        
        # ดึง room_id
        if self.admin_comein == True:
            room_id = self.view.get_type_search()
        else:
            room_id = self.log_room_id if hasattr(self, 'log_room_id') else None
        
        # print(f"DEBUG: room_id = {room_id}")
        
        if room_id is None:
            # print("DEBUG: room_id เป็น None")
            return
        
        # ดึงรายการ templates จาก database
        templates = self.receive_lab_service.get_report_templates(room_id)
        
        # เก็บไว้ใน all_templates เพื่อใช้ตอน export
        self.all_templates = templates
        
        # print(f"DEBUG: พบ template {len(templates)} รายการจาก database")
        # for tmpl in templates:
        #     print(f"  - {tmpl.get('report_name', '')}")
        
        if not templates:
            # print("DEBUG: ไม่มี templates จาก database")
            return
        
        # ค้นหา template ที่ตรงกับรายการตรวจ
        matched_templates = set()  # ใช้ set เพื่อไม่ให้ซ้ำ
        
        # print(f"\nDEBUG: กำลังค้นหาจาก test_items จำนวน {len(test_items)} รายการ:")
        for test_item in test_items:
            test_name = test_item.get('test_name', '').strip()
            # print(f"  - ชื่อรายการตรวจ: '{test_name}'")
            
            if not test_name:
                continue
            
            # ค้นหา template ที่มีชื่อตรงกับรายการตรวจ
            for template in templates:
                template_name = template.get('report_name', '')
                # ตรวจสอบว่า test_name อยู่ใน template_name หรือไม่
                if test_name.lower() in template_name.lower():
                    # print(f"    ✓ ตรงกับ: {template_name}")
                    matched_templates.add(template_name)
        
        # print(f"\nDEBUG: พบ template ที่ตรงกัน {len(matched_templates)} รายการ")
        
        # แสดงผลลัพธ์ใน tableView_3
        if matched_templates:
            for template_name in sorted(matched_templates):
                name_item = QStandardItem(template_name)
                self.template_model.appendRow([name_item])
        else:
            # ไม่พบ template ที่ตรงกัน
            no_template_item = QStandardItem("ไม่พบ Template ที่ตรงกัน")
            self.template_model.appendRow([no_template_item])
    
    # ==========================================
    # LAB ORDER ACTIONS - RECEIVE & REJECT - การรับและปฏิเสธแลป
    # ==========================================
    def receive_lab_orders(self):
        """บันทึกการรับแลป"""
        # ตรวจสอบว่ามีการเลือก Lab Order หรือไม่
        if not hasattr(self, 'current_lab_order_id') or self.current_lab_order_id is None:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณาเลือก Lab Order ที่ต้องการรับก่อน (ดับเบิ้ลคลิกที่รายการ)")
            return
        
        # ตรวจสอบว่าเลือกสถานะตัวอย่างหรือไม่
        sample_status = ""
        if self.view.ui.sample_status_good_radioButton.isChecked():
            sample_status = "ปกติ"
        elif self.view.ui.sample_status_bad_radioButton.isChecked():
            sample_status = "เสียหาย/ไม่ปกติ"
        else:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณาเลือกสภาพสิ่งส่งตรวจ (ปกติ หรือ เสียหาย/ไม่ปกติ)")
            return
        
        # ดึง comment
        comment = self.view.ui.comment_status_textEdit.toPlainText().strip()
        
        # ดึง room_id
        room_id = self.log_room_id if hasattr(self, 'log_room_id') else None
        if room_id is None:
            QMessageBox.warning(self.view, "ข้อผิดพลาด", "ไม่พบข้อมูลห้องแลป")
            return
        
        # ดึง employee_id จาก main_controller
        employee_id = None
        if self.main_controller:
            employee_id = self.main_controller.get_user_login_id()
        
        if employee_id is None:
            QMessageBox.warning(self.view, "ข้อผิดพลาด", "ไม่พบข้อมูลผู้ใช้งาน")
            return
        
        try:
            # เรียก API เพื่อบันทึกการรับแลป
            result = self.receive_lab_service.receive_lab_order(
                lab_order_id=self.current_lab_order_id,
                receive_from_room=room_id,
                comment_for_sample=comment,
                sample_status=sample_status,
                updater_id=employee_id
            )
            
            if result and result.get('success', False):
                QMessageBox.information(self.view, "สำเร็จ", result.get('message', 'บันทึกการรับแลปสำเร็จ'))
                
                # แสดงไฟล์ template หลังจากรับแลปสำเร็จ
                if self.current_test_items:
                    self.find_and_display_matching_templates(self.current_test_items)
                
                # เคลียร์ข้อมูล
                self.view.ui.comment_status_textEdit.clear()
                self.view.ui.sample_status_good_radioButton.setAutoExclusive(False)
                self.view.ui.sample_status_bad_radioButton.setAutoExclusive(False)
                self.view.ui.sample_status_good_radioButton.setChecked(False)
                self.view.ui.sample_status_bad_radioButton.setChecked(False)
                self.view.ui.sample_status_good_radioButton.setAutoExclusive(True)
                self.view.ui.sample_status_bad_radioButton.setAutoExclusive(True)
                # หมายเหตุ: ไม่เคลียร์ current_lab_order_id เพื่อให้สามารถ Export ได้
                # self.current_lab_order_id = None
            else:
                message = result.get('message', 'ไม่สามารถบันทึกได้') if result else 'ไม่สามารถบันทึกได้'
                QMessageBox.warning(self.view, "ข้อผิดพลาด", message)
                
        except Exception as e:
            QMessageBox.critical(self.view, "ข้อผิดพลาด", f"เกิดข้อผิดพลาด: {str(e)}")
        
    
    def reject_lab_orders(self):
        # ตรวจสอบว่ามีการเลือก Lab Order หรือไม่
        if self.current_lab_order_id is None:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณาดับเบิลคลิกเลือก Lab Order ที่ต้องการปฏิเสธ")
            return
        
        # ตรวจสอบสภาพสิ่งส่งตรวจ
        if self.view.ui.sample_status_good_radioButton.isChecked():
            sample_status = "ปกติ"
        elif self.view.ui.sample_status_bad_radioButton.isChecked():
            sample_status = "เสียหาย/ไม่ปกติ"
        else:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณาเลือกสภาพสิ่งส่งตรวจ (ปกติ หรือ เสียหาย/ไม่ปกติ)")
            return
        
        # ดึง comment
        comment = self.view.ui.comment_status_textEdit.toPlainText().strip()
        
        # ดึง room_id
        room_id = self.log_room_id if hasattr(self, 'log_room_id') else None
        if room_id is None:
            QMessageBox.warning(self.view, "ข้อผิดพลาด", "ไม่พบข้อมูลห้องแลป")
            return
        
        # ดึง employee_id จาก main_controller
        employee_id = None
        if self.main_controller:
            employee_id = self.main_controller.get_user_login_id()
        
        if employee_id is None:
            QMessageBox.warning(self.view, "ข้อผิดพลาด", "ไม่พบข้อมูลผู้ใช้งาน")
            return
        
        try:
            # เรียก API เพื่อบันทึกการปฏิเสธแลป
            result = self.receive_lab_service.reject_lab_order(
                lab_order_id=self.current_lab_order_id,
                receive_from_room=room_id,
                comment_for_sample=comment,
                sample_status=sample_status,
                updater_id=employee_id
            )
            
            if result and result.get('success', False):
                QMessageBox.information(self.view, "สำเร็จ", result.get('message', 'บันทึกการปฏิเสธแลปสำเร็จ'))
                # เคลียร์ข้อมูล
                self.view.ui.comment_status_textEdit.clear()
                self.view.ui.sample_status_good_radioButton.setAutoExclusive(False)
                self.view.ui.sample_status_bad_radioButton.setAutoExclusive(False)
                self.view.ui.sample_status_good_radioButton.setChecked(False)
                self.view.ui.sample_status_bad_radioButton.setChecked(False)
                self.view.ui.sample_status_good_radioButton.setAutoExclusive(True)
                self.view.ui.sample_status_bad_radioButton.setAutoExclusive(True)
                self.current_lab_order_id = None
            else:
                message = result.get('message', 'ไม่สามารถบันทึกได้') if result else 'ไม่สามารถบันทึกได้'
                QMessageBox.warning(self.view, "ข้อผิดพลาด", message)
                
        except Exception as e:
            QMessageBox.critical(self.view, "ข้อผิดพลาด", f"เกิดข้อผิดพลาด: {str(e)}")

    # ==========================================
    # BUTTON HANDLERS - EXPORT & CLEAR - จัดการปุ่ม Export และ Clear
    # ==========================================
    def export_pushButton_clicked(self):
        """เมื่อกด Export จะคัดลอกไฟล์ template พร้อมเติมข้อมูลไปยัง location ที่กำหนด"""
        # ตรวจสอบว่ามีการเลือก template หรือไม่
        if not self.selected_template:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณาเลือกไฟล์ Template ที่ต้องการ Export ก่อน")
            return
        
        # ตรวจสอบว่ามีการเลือก Lab Order หรือไม่
        if not hasattr(self, 'current_lab_order_id') or self.current_lab_order_id is None:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณาเลือก Lab Order ก่อน")
            return
        
        template_name = self.selected_template.get('report_name', '')
        template_path = self.selected_template.get('report_path', '')
        
        print(f"DEBUG Export: template_name = {template_name}")
        print(f"DEBUG Export: template_path = {template_path}")
        
        # สร้าง full path ของไฟล์
        import os
        source_file = os.path.join(template_path, template_name)
        print(f"DEBUG Export: source_file = {source_file}")
        
        # ตรวจสอบว่าไฟล์มีอยู่จริงหรือไม่
        if not os.path.exists(source_file):
            QMessageBox.warning(self.view, "ข้อผิดพลาด", f"ไม่พบไฟล์ Template:\n{source_file}")
            return
        
        # เปิด file dialog ให้เลือกสถานที่บันทึก
        from PySide6.QtWidgets import QFileDialog
        
        # ดึงนามสกุลไฟล์
        file_extension = os.path.splitext(template_name)[1]
        filter_text = f"Word Documents (*{file_extension})" if file_extension == '.docx' else f"All Files (*{file_extension})"
        
        # สร้างชื่อไฟล์จาก lab_order_id + ส่วนสำคัญจาก template name
        lab_order_str = str(self.current_lab_order_id).zfill(12)
        
        # ตัดเอาส่วนสำคัญจากชื่อ template
        # เช่น "2.Template ห้องปฏิบัติการแบคทีเรียวิทยา Ver Blind_VITEK2 with MIC.docx"
        # -> "VITEK2 with MIC"
        template_name_without_ext = os.path.splitext(template_name)[0]
        
        # หาส่วนที่มี "VITEK" หรือคำสำคัญอื่นๆ
        template_suffix = ""
        if "VITEK2 with MIC" in template_name_without_ext:
            template_suffix = "_VITEK2 with MIC"
        elif "VITEK2 iden" in template_name_without_ext:
            template_suffix = "_VITEK2 iden"
        elif "MIC" in template_name_without_ext and "VITEK2" not in template_name_without_ext:
            template_suffix = "_MIC"
        # เพิ่มเงื่อนไขอื่นๆ ตามต้องการ
        
        default_filename = f"{lab_order_str}{template_suffix}{file_extension}"
        
        save_path, _ = QFileDialog.getSaveFileName(
            self.view,
            "บันทึกไฟล์ Template",
            default_filename,
            filter_text
        )
        
        if not save_path:
            print("DEBUG Export: ยกเลิกการบันทึก")
            return
        
        # ดึงข้อมูลจาก database และเติมลงใน template
        try:
            data = self.get_template_data()
            if data:
                self.fill_word_template(source_file, save_path, data)
                print(f"DEBUG Export: เติมข้อมูลและบันทึกไฟล์สำเร็จ -> {save_path}")
                QMessageBox.information(self.view, "สำเร็จ", f"บันทึกไฟล์สำเร็จ:\n{save_path}")
            else:
                QMessageBox.warning(self.view, "ข้อผิดพลาด", "ไม่สามารถดึงข้อมูลจาก database ได้")
        except Exception as e:
            print(f"DEBUG Export Error: {e}")
            import traceback
            print(traceback.format_exc())
            QMessageBox.critical(self.view, "ข้อผิดพลาด", f"ไม่สามารถเติมข้อมูลและบันทึกไฟล์ได้:\n{str(e)}")
    
    def get_template_data(self):
        """ดึงข้อมูลจาก database สำหรับเติมลงใน template"""
        try:
            import sys
            import os
            from datetime import datetime
            
            print(f"DEBUG get_template_data: เริ่มดึงข้อมูลสำหรับ lab_order_id = {self.current_lab_order_id}")
            
            # เพิ่ม BACKEND path
            current_file = os.path.abspath(__file__)
            report_lab_path = os.path.dirname(os.path.dirname(current_file))
            project_path = os.path.dirname(report_lab_path)
            backend_path = os.path.join(project_path, 'BACKEND')
            
            if backend_path not in sys.path:
                sys.path.insert(0, backend_path)

            from BACKEND.database import get_db_connection  # type: ignore

            conn = get_db_connection()
            if conn is None:
                print("DEBUG get_template_data: ไม่สามารถเชื่อมต่อ database")
                return None
            
            cursor = conn.cursor()
            
            # 1. ดึง id จาก lab_receive_detail (เลขที่รายงาน)
            query1 = """
                SELECT id, dtime
                FROM lab_receive_detail 
                WHERE lab_order_id = ? 
                ORDER BY id DESC 
                LIMIT 1
            """
            print(f"DEBUG: Query1 - lab_order_id = {self.current_lab_order_id}")
            cursor.execute(query1, (self.current_lab_order_id,))
            result1 = cursor.fetchone()
            report_id = result1[0] if result1 else None
            receive_dtime = result1[1] if result1 else None
            print(f"DEBUG: ผลลัพธ์ Query1 - report_id = {report_id}, receive_dtime = {receive_dtime}")
            
            # 2. ดึงข้อมูลจาก lab_order และ sample_registration
            query2 = """
                SELECT sr.collect_date, sr.id as case_id, sr.name, sr.species, sr.breed, 
                       sr.sex, sr.age_year, sr.age_month, sr.age_day, sr.sample_type, 
                       lo.id as lab_order_id, sr.case_id
                FROM lab_order lo
                LEFT JOIN sample_registration sr ON lo.sample_id = sr.id
                WHERE lo.id = ?
            """
            print(f"DEBUG: Query2 - lab_order_id = {self.current_lab_order_id}")
            cursor.execute(query2, (self.current_lab_order_id,))
            result2 = cursor.fetchone()
            
            if result2:
                collect_date = result2[0]
                case_id = result2[1]
                animal_name = result2[2]
                species = result2[3]
                breed = result2[4]
                sex = result2[5]
                age_year = result2[6]
                age_month = result2[7]
                age_day = result2[8]
                sample_type = result2[9]
                registration_case_id = result2[11]  # case_id จาก sample_registration
                print(f"DEBUG: ผลลัพธ์ Query2 - collect_date = {collect_date}, case_id = {case_id}")
                print(f"DEBUG: animal_name = {animal_name}, species = {species}, breed = {breed}")
                print(f"DEBUG: age_year = {age_year}, age_month = {age_month}, age_day = {age_day}")
                print(f"DEBUG: registration_case_id = {registration_case_id}")
            else:
                print("DEBUG: ไม่พบข้อมูลจาก Query2")
                collect_date = case_id = animal_name = species = breed = sex = None
                age_year = age_month = age_day = sample_type = registration_case_id = None
            
            # 3. ดึงข้อมูล owner และ sender จาก case_registration และ customer
            owner_name = owner_phone = owner_email = owner_address = ''
            sender_name = sender_phone = sender_email = sender_address = ''
            
            if registration_case_id:
                query3 = """
                    SELECT cr.owner_id, cr.sender_id,
                           o.name as owner_name, o.surname as owner_surname, 
                           o.phone as owner_phone, o.email as owner_email, 
                           o.contact_address as owner_address,
                           s.name as sender_name, s.surname as sender_surname,
                           s.phone as sender_phone, s.email as sender_email,
                           s.contact_address as sender_address
                    FROM case_registration cr
                    LEFT JOIN customer o ON cr.owner_id = o.id
                    LEFT JOIN customer s ON cr.sender_id = s.id
                    WHERE cr.id = ?
                """
                print(f"DEBUG: Query3 - registration_case_id = {registration_case_id}")
                cursor.execute(query3, (registration_case_id,))
                result3 = cursor.fetchone()
                
                if result3:
                    owner_name = f"{result3[2] or ''} {result3[3] or ''}".strip()
                    owner_phone = result3[4] or ''
                    owner_email = result3[5] or ''
                    owner_address = result3[6] or ''
                    
                    sender_name = f"{result3[7] or ''} {result3[8] or ''}".strip()
                    sender_phone = result3[9] or ''
                    sender_email = result3[10] or ''
                    sender_address = result3[11] or ''
                    
                    print(f"DEBUG: owner_name = {owner_name}, owner_phone = {owner_phone}")
                    print(f"DEBUG: sender_name = {sender_name}, sender_phone = {sender_phone}")
                else:
                    print("DEBUG: ไม่พบข้อมูลจาก Query3")
            
            cursor.close()
            conn.close()
            
            # 4. สร้างเลขที่ตัวอย่าง: D(day today)-lab_order_id
            today = datetime.now()
            day_str = str(today.day)  # เอาเฉพาะวัน
            sample_number = f"D{day_str}-{self.current_lab_order_id}"
            
            # 5. จัดรูปแบบอายุ
            age_text = ''
            age_parts = []
            if age_year:
                age_parts.append(f"{age_year} ปี")
            if age_month:
                age_parts.append(f"{age_month} เดือน")
            if age_day:
                age_parts.append(f"{age_day} วัน")
            age_text = ' '.join(age_parts)
            
            # 6. แปลงวันที่ให้เป็นรูปแบบ DD/MM/YYYY
            formatted_collect_date = ''
            if collect_date:
                try:
                    # ลอง parse วันที่จาก string
                    if isinstance(collect_date, str):
                        # ถ้าเป็น format YYYY-MM-DD หรือ YYYY-MM-DD HH:MM:SS
                        date_part = collect_date.split(' ')[0]  # เอาเฉพาะส่วนวันที่
                        from datetime import datetime as dt
                        date_obj = dt.strptime(date_part, '%Y-%m-%d')
                        formatted_collect_date = date_obj.strftime('%d/%m/%Y')
                    else:
                        formatted_collect_date = str(collect_date)
                except:
                    formatted_collect_date = str(collect_date)
            
            data = {
                'report_id': str(report_id) if report_id else '',
                'sample_number': sample_number,
                'collect_date': formatted_collect_date,
                'lab_order_id': self.current_lab_order_id,
                'case_id': case_id,
                'animal_name': animal_name if animal_name else '',
                'species': species if species else '',
                'breed': breed if breed else '',
                'sex': sex if sex else '',
                'age': age_text,
                'sample_type': sample_type if sample_type else '',
                'receive_dtime': str(receive_dtime) if receive_dtime else '',
                'owner_name': owner_name,
                'owner_phone': owner_phone,
                'owner_email': owner_email,
                'owner_address': owner_address,
                'sender_name': sender_name,
                'sender_phone': sender_phone,
                'sender_email': sender_email,
                'sender_address': sender_address
            }
            
            print(f"DEBUG get_template_data: ข้อมูลที่ได้ = {data}")
            return data
            
        except Exception as e:
            print(f"ERROR in get_template_data: {e}")
            import traceback
            print(traceback.format_exc())
            return None
    
    def fill_word_template(self, source_file, output_file, data):
        """เติมข้อมูลลงใน Word template"""
        try:
            from docx import Document
            
            print(f"DEBUG fill_word_template: เริ่มเติมข้อมูล")
            print(f"DEBUG: source_file = {source_file}")
            print(f"DEBUG: output_file = {output_file}")
            
            # เปิดไฟล์ template
            doc = Document(source_file)
            print(f"DEBUG: เปิดไฟล์ template สำเร็จ")
            
            # สร้าง mapping ของข้อมูล (แทนที่เฉพาะส่วนหลัง colon)
            # หมายเหตุ: ฟิลด์ที่อยู่ใน table และจัดการแยก (next_cell) ไม่ควรอยู่ใน replacements นี้
            replacements = {
                'เลขที่รายงาน:': f"เลขที่รายงาน: {data.get('report_id', '')}",
                'วันที่รับตัวอย่าง:': f"วันที่รับตัวอย่าง: {data.get('collect_date', '')}",
                'เลขที่ตัวอย่าง:': f"เลขที่ตัวอย่าง: {data.get('sample_number', '')}",
                'พันธุ์:': f"พันธุ์: {data.get('breed', '')}",
                'เพศ:': f"เพศ: {data.get('sex', '')}",
                # ฟิลด์อื่นๆ (ชื่อเจ้าของ, ชื่อผู้ส่ง, ชนิดสัตว์, ชื่อสัตว์, อายุ, ชนิดตัวอย่าง) 
                # ถูกจัดการแยกด้วย next_cell logic ด้านล่าง
            }
            
            # เพิ่ม mapping สำหรับ โทร. และ E-mail: ของเจ้าของและผู้ส่ง
            # ใช้ startswith เพราะอาจมีหลายแบบ
            phone_email_replacements = {
                ('ชื่อเจ้าของ:', 'โทร.'): data.get('owner_phone', ''),
                ('ชื่อเจ้าของ:', 'E-mail:'): data.get('owner_email', ''),
                ('ที่อยู่:', 'owner'): data.get('owner_address', ''),
                ('ชื่อผู้ส่ง:', 'โทร.'): data.get('sender_phone', ''),
                ('ชื่อผู้ส่ง:', 'E-mail:'): data.get('sender_email', ''),
                ('ที่อยู่:', 'sender'): data.get('sender_address', '')
            }
            
            print(f"DEBUG fill_word_template: replacements = {replacements}")
            print(f"DEBUG: จำนวน paragraphs = {len(doc.paragraphs)}")
            print(f"DEBUG: จำนวน tables = {len(doc.tables)}")
            
            # นับจำนวนการแทนที่
            replacement_count = 0
            
            # ฟังก์ชันแทนที่ข้อความใน paragraph
            def replace_text_in_paragraph(paragraph):
                nonlocal replacement_count
                full_text = paragraph.text
                
                # ตรวจสอบและแทนที่
                for key, value in replacements.items():
                    if key in full_text:
                        print(f"DEBUG: พบคำว่า '{key}' ใน paragraph -> จะแทนที่ด้วย '{value}'")
                        
                        # เก็บ formatting จาก run แรก
                        original_font = None
                        if len(paragraph.runs) > 0:
                            original_run = paragraph.runs[0]
                            original_font = original_run.font
                        
                        # ลบ runs ทั้งหมด
                        inline = paragraph.runs
                        for i in range(len(inline)-1, -1, -1):
                            p = inline[i]._element
                            p.getparent().remove(p)
                        
                        # สร้าง run ใหม่พร้อมข้อความที่แทนที่แล้ว
                        new_text = full_text.replace(key, value)
                        new_run = paragraph.add_run(new_text)
                        
                        # คัดลอก formatting กลับมา
                        if original_font:
                            new_run.font.name = original_font.name
                            new_run.font.size = original_font.size
                            new_run.font.bold = original_font.bold
                            new_run.font.italic = original_font.italic
                            new_run.font.underline = original_font.underline
                            new_run.font.color.rgb = original_font.color.rgb
                        
                        replacement_count += 1
                        print(f"DEBUG: แทนที่สำเร็จ -> '{new_text[:50]}...'")
                        break  # หยุดหลังแทนที่ครั้งแรก
            
            # แทนที่ข้อความใน paragraphs
            print(f"DEBUG: กำลังตรวจสอบ paragraphs...")
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # เฉพาะ paragraph ที่มีข้อความ
                    print(f"DEBUG: Paragraph {i}: {paragraph.text[:100]}")
                replace_text_in_paragraph(paragraph)
            
            # แทนที่ข้อความใน tables
            print(f"DEBUG: กำลังตรวจสอบ tables...")
            processed_cells = set()  # เก็บ (row_idx, cell_idx) ของ cell ที่ประมวลผลแล้ว
            
            for table_idx, table in enumerate(doc.tables):
                print(f"DEBUG: Table {table_idx}")
                for row_idx, row in enumerate(table.rows):
                    # ตรวจสอบว่าแถวนี้มี "ชื่อเจ้าของ:" หรือ "ชื่อผู้ส่ง:" หรือไม่
                    row_text = ' '.join([cell.text for cell in row.cells])
                    
                    for cell_idx, cell in enumerate(row.cells):
                        # ตรวจสอบ Row 4 (ที่อยู่เจ้าของ) และ Row 6 (ที่อยู่ผู้ส่ง) ก่อนเช็ค merged cell
                        if row_idx == 4 and cell_idx == 0:
                            # Row 4 = ที่อยู่ของเจ้าของ
                            print(f"DEBUG: Row 4 Cell[0] - ใส่ที่อยู่เจ้าของ: {data.get('owner_address', '')}")
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.text = ''
                                if len(para.runs) > 0:
                                    para.runs[0].text = f"ที่อยู่: {data.get('owner_address', '')}"
                                else:
                                    para.add_run(f"ที่อยู่: {data.get('owner_address', '')}")
                                replacement_count += 1
                                print(f"DEBUG: ใส่ owner_address สำเร็จ")
                                break
                            continue  # ข้ามการตรวจสอบอื่นๆ
                        
                        if row_idx == 6 and cell_idx == 0:
                            # Row 6 = ที่อยู่ของผู้ส่ง
                            print(f"DEBUG: Row 6 Cell[0] - ใส่ที่อยู่ผู้ส่ง: {data.get('sender_address', '')}")
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.text = ''
                                if len(para.runs) > 0:
                                    para.runs[0].text = f"ที่อยู่: {data.get('sender_address', '')}"
                                else:
                                    para.add_run(f"ที่อยู่: {data.get('sender_address', '')}")
                                replacement_count += 1
                                print(f"DEBUG: ใส่ sender_address สำเร็จ")
                                break
                            continue  # ข้ามการตรวจสอบอื่นๆ
                        
                        # ใช้ (row_idx, cell_idx) เพื่อ identify cell แทน id(cell._element)
                        cell_position = (row_idx, cell_idx)
                        
                        # แสดง cell ก่อนตรวจสอบ
                        cell_text_preview = cell.text.strip()[:50] if cell.text.strip() else '(empty)'
                        
                        # ตรวจสอบว่า cell นี้เป็น merged cell หรือไม่
                        cell_id = id(cell._element)
                        is_merged = False
                        for prev_cell_idx in range(cell_idx):
                            if id(row.cells[prev_cell_idx]._element) == cell_id:
                                is_merged = True
                                break
                        
                        if is_merged:
                            print(f"DEBUG: SKIPPED Table[{table_idx}] Row[{row_idx}] Cell[{cell_idx}]: {cell_text_preview} (merged cell)")
                            continue
                        
                        if cell_position in processed_cells:
                            print(f"DEBUG: SKIPPED Table[{table_idx}] Row[{row_idx}] Cell[{cell_idx}]: {cell_text_preview} (already processed)")
                            continue
                        
                        processed_cells.add(cell_position)
                        
                        # ตรวจสอบทุก paragraph ใน cell
                        for para in cell.paragraphs:
                            cell_text = para.text.strip()
                            if cell_text:
                                print(f"DEBUG: Table[{table_idx}] Row[{row_idx}] Cell[{cell_idx}]: {cell_text[:50]}")
                            
                            # จัดการฟิลด์พิเศษ: ชื่อเจ้าของ, โทร., E-mail:, ที่อยู่:
                            if 'ชื่อเจ้าของ:' in row_text or 'ชื่อ:' in row_text:
                                # แถวนี้เป็นของเจ้าของ
                                print(f"DEBUG: เจอแถวเจ้าของ - cell_text = '{cell_text}'")
                                if cell_text.startswith('ชื่อเจ้าของ:') and cell_idx == 0:
                                    # Row 3 Cell[0-3] เป็น merged cell, ใส่ข้อมูลใน cell เดียวกัน
                                    print(f"DEBUG: เจอ label ชื่อเจ้าของ - จะใส่ข้อมูลใน cell เดียวกัน")
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.text = ''
                                        if len(para.runs) > 0:
                                            para.runs[0].text = f"ชื่อเจ้าของ: {data.get('owner_name', '')}"
                                        else:
                                            para.add_run(f"ชื่อเจ้าของ: {data.get('owner_name', '')}")
                                        replacement_count += 1
                                        print(f"DEBUG: ใส่ owner_name: {data.get('owner_name', '')} สำเร็จ")
                                        break
                                elif cell_text == 'โทร.' and cell_idx == 4:
                                    # Row 3 Cell[4-5] เป็น merged cell, ใส่ข้อมูลใน cell เดียวกัน
                                    print(f"DEBUG: เจอ โทร. เจ้าของ - จะใส่ '{data.get('owner_phone', '')}' ใน cell เดียวกัน")
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.text = ''
                                        if len(para.runs) > 0:
                                            para.runs[0].text = f"โทร. {data.get('owner_phone', '')}"
                                        else:
                                            para.add_run(f"โทร. {data.get('owner_phone', '')}")
                                        replacement_count += 1
                                        print(f"DEBUG: ใส่ owner_phone: {data.get('owner_phone', '')} สำเร็จ")
                                        break
                                elif cell_text == 'E-mail:' and cell_idx == 6:
                                    # Row 3 Cell[6-7] เป็น merged cell, ใส่ข้อมูลใน cell เดียวกัน
                                    print(f"DEBUG: เจอ E-mail เจ้าของ - จะใส่ '{data.get('owner_email', '')}' ใน cell เดียวกัน")
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.text = ''
                                        if len(para.runs) > 0:
                                            para.runs[0].text = f"E-mail: {data.get('owner_email', '')}"
                                        else:
                                            para.add_run(f"E-mail: {data.get('owner_email', '')}")
                                        replacement_count += 1
                                        print(f"DEBUG: ใส่ owner_email: {data.get('owner_email', '')} สำเร็จ")
                                        break
                                elif cell_text.startswith('ที่อยู่:') and row_idx == 4:
                                    # Row 4 Cell[0-7] เป็น merged cell ทั้งหมด, ใส่ข้อมูลใน cell เดียวกัน
                                    print(f"DEBUG: เจอ ที่อยู่ เจ้าของ (row_idx=4) - จะใส่ '{data.get('owner_address', '')}' ใน cell เดียวกัน")
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.text = ''
                                        if len(para.runs) > 0:
                                            para.runs[0].text = f"ที่อยู่: {data.get('owner_address', '')}"
                                        else:
                                            para.add_run(f"ที่อยู่: {data.get('owner_address', '')}")
                                        replacement_count += 1
                                        print(f"DEBUG: ใส่ owner_address: {data.get('owner_address', '')} สำเร็จ")
                                        break
                            
                            elif 'ชื่อผู้ส่ง:' in row_text:
                                # แถวนี้เป็นของผู้ส่ง
                                if cell_text.startswith('ชื่อผู้ส่ง:') and cell_idx == 0:
                                    # Row 5 Cell[0-3] เป็น merged cell, ใส่ข้อมูลใน cell เดียวกัน
                                    print(f"DEBUG: เจอ label ชื่อผู้ส่ง - จะใส่ข้อมูลใน cell เดียวกัน")
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.text = ''
                                        if len(para.runs) > 0:
                                            para.runs[0].text = f"ชื่อผู้ส่ง: {data.get('sender_name', '')}"
                                        else:
                                            para.add_run(f"ชื่อผู้ส่ง: {data.get('sender_name', '')}")
                                        replacement_count += 1
                                        print(f"DEBUG: ใส่ sender_name: {data.get('sender_name', '')} สำเร็จ")
                                        break
                                elif cell_text == 'โทร.' and cell_idx == 4:
                                    # Row 5 Cell[4-5] เป็น merged cell, ใส่ข้อมูลใน cell เดียวกัน
                                    print(f"DEBUG: เจอ โทร. ผู้ส่ง - จะใส่ '{data.get('sender_phone', '')}' ใน cell เดียวกัน")
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.text = ''
                                        if len(para.runs) > 0:
                                            para.runs[0].text = f"โทร. {data.get('sender_phone', '')}"
                                        else:
                                            para.add_run(f"โทร. {data.get('sender_phone', '')}")
                                        replacement_count += 1
                                        print(f"DEBUG: ใส่ sender_phone: {data.get('sender_phone', '')} สำเร็จ")
                                        break
                                elif cell_text == 'E-mail:' and cell_idx == 6:
                                    # Row 5 Cell[6-7] เป็น merged cell, ใส่ข้อมูลใน cell เดียวกัน
                                    print(f"DEBUG: เจอ E-mail ผู้ส่ง - จะใส่ '{data.get('sender_email', '')}' ใน cell เดียวกัน")
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.text = ''
                                        if len(para.runs) > 0:
                                            para.runs[0].text = f"E-mail: {data.get('sender_email', '')}"
                                        else:
                                            para.add_run(f"E-mail: {data.get('sender_email', '')}")
                                        replacement_count += 1
                                        print(f"DEBUG: ใส่ sender_email: {data.get('sender_email', '')} สำเร็จ")
                                        break
                                elif cell_text.startswith('ที่อยู่:') and row_idx == 6:
                                    # Row 6 Cell[0-7] เป็น merged cell ทั้งหมด, ใส่ข้อมูลใน cell เดียวกัน
                                    print(f"DEBUG: เจอ ที่อยู่ ผู้ส่ง (row_idx=6) - จะใส่ '{data.get('sender_address', '')}' ใน cell เดียวกัน")
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.text = ''
                                        if len(para.runs) > 0:
                                            para.runs[0].text = f"ที่อยู่: {data.get('sender_address', '')}"
                                        else:
                                            para.add_run(f"ที่อยู่: {data.get('sender_address', '')}")
                                        replacement_count += 1
                                        print(f"DEBUG: ใส่ sender_address: {data.get('sender_address', '')} สำเร็จ")
                                        break
                            
                            # จัดการฟิลด์สัตว์: ชนิดสัตว์, ชื่อสัตว์, ชนิดตัวอย่าง, อายุ
                            print(f"DEBUG: ตรวจสอบฟิลด์สัตว์ - cell_text = '{cell_text}'")
                            if cell_text.startswith('ชนิดสัตว์:'):
                                # ROW 7 Cell[0] - ไม่มี cell ถัดไป (Cell[1] คือ label "ชื่อสัตว์:")
                                # ต้องใส่ข้อมูลใน cell เดียวกัน
                                print(f"DEBUG: เจอ ชนิดสัตว์ - จะใส่ '{data.get('species', '')}' ใน cell เดียวกัน")
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ''
                                    if len(para.runs) > 0:
                                        para.runs[0].text = f"ชนิดสัตว์: {data.get('species', '')}"
                                    else:
                                        para.add_run(f"ชนิดสัตว์: {data.get('species', '')}")
                                    replacement_count += 1
                                    print(f"DEBUG: ใส่ species: {data.get('species', '')} สำเร็จ")
                                    break
                            elif cell_text.startswith('ชื่อสัตว์:') and cell_idx == 1:
                                # ROW 7 Cell[1] - label, ใส่ข้อมูลใน Cell[2-4] (merged)
                                print(f"DEBUG: เจอ ชื่อสัตว์ label - จะใส่ '{data.get('animal_name', '')}' ใน cell[2]")
                                if cell_idx + 1 < len(row.cells):
                                    next_cell = row.cells[cell_idx + 1]  # Cell[2]
                                    for next_para in next_cell.paragraphs:
                                        for run in next_para.runs:
                                            run.text = ''
                                        if len(next_para.runs) > 0:
                                            next_para.runs[0].text = data.get('animal_name', '')
                                        else:
                                            next_para.add_run(data.get('animal_name', ''))
                                        replacement_count += 1
                                        print(f"DEBUG: ใส่ animal_name: {data.get('animal_name', '')} สำเร็จ")
                                        break
                            elif cell_text.startswith('ชนิดตัวอย่าง:') and cell_idx == 0:
                                # ROW 8 Cell[0-6] - merged cell, ใส่ข้อมูลใน cell เดียวกัน
                                print(f"DEBUG: เจอ ชนิดตัวอย่าง - จะใส่ '{data.get('sample_type', '')}' ใน cell เดียวกัน")
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ''
                                    if len(para.runs) > 0:
                                        para.runs[0].text = f"ชนิดตัวอย่าง: {data.get('sample_type', '')}"
                                    else:
                                        para.add_run(f"ชนิดตัวอย่าง: {data.get('sample_type', '')}")
                                    replacement_count += 1
                                    print(f"DEBUG: ใส่ sample_type: {data.get('sample_type', '')} สำเร็จ")
                                    break
                            elif cell_text.startswith('อายุ:') and cell_idx == 7:
                                # ROW 8 Cell[7] - cell สุดท้าย ไม่มี cell ถัดไป, ใส่ใน cell เดียวกัน
                                print(f"DEBUG: เจอ อายุ - จะใส่ '{data.get('age', '')}' ใน cell เดียวกัน")
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ''
                                    if len(para.runs) > 0:
                                        para.runs[0].text = f"อายุ: {data.get('age', '')}"
                                    else:
                                        para.add_run(f"อายุ: {data.get('age', '')}")
                                    replacement_count += 1
                                    print(f"DEBUG: ใส่ age: {data.get('age', '')} สำเร็จ")
                                    break
                            
                            # แทนที่ข้อความปกติ
                            replace_text_in_paragraph(para)
            
            print(f"DEBUG: แทนที่ข้อความทั้งหมด {replacement_count} ครั้ง")
            
            # บันทึกไฟล์
            doc.save(output_file)
            print(f"DEBUG fill_word_template: บันทึกไฟล์สำเร็จ -> {output_file}")
            
        except Exception as e:
            print(f"ERROR in fill_word_template: {e}")
            import traceback
            print(traceback.format_exc())
            raise

    def clear_pushButton_clicked(self):
        self.view.ui.barcode_lineEdit.clear()
        self.model.removeRows(0, self.model.rowCount())
        self.detail_model.removeRows(0, self.detail_model.rowCount())
        self.template_model.removeRows(0, self.template_model.rowCount())
        self.reset_lazy_loading_state()
    
    # ==========================================
    # DATA MANAGEMENT - RESET & CLEAR - จัดการข้อมูล รีเซ็ตและเคลียร์
    # ==========================================
    def reset_lazy_loading_state(self):
        self.current_offset = 0
        self.is_loading = False
        self.has_more_data = True
        self.all_data = []
    
    def clear_all_data(self):
        """เคลียร์ข้อมูลทั้งหมดเมื่อ login เข้ามาใหม่"""
        self.view.ui.barcode_lineEdit.clear()
        self.model.removeRows(0, self.model.rowCount())
        self.detail_model.removeRows(0, self.detail_model.rowCount())
        self.template_model.removeRows(0, self.template_model.rowCount())
        self.reset_lazy_loading_state()
    
    # ==========================================
    # LAZY LOADING - SCROLL HANDLER - โหลดข้อมูลแบบค่อยเป็นค่อยไป
    # ==========================================
    def on_scroll(self, value):
        scrollbar = self.view.ui.tableView.verticalScrollBar()
        if value >= scrollbar.maximum() - 10:
            if self.has_more_data and not self.is_loading:
                barcode = self.view.ui.barcode_lineEdit.text()
                if barcode == "":
                    self.load_more_lab_orders()
    
    # ==========================================
    # DATA LOADING - SEARCH & LOAD - โหลดข้อมูล ค้นหาและโหลด
    # ==========================================
    def loaded_lab_orders(self):
        barcode = self.view.ui.barcode_lineEdit.text().strip()
        if barcode == "":
            self.view.clear_all_table()
            self.model.removeRows(0, self.model.rowCount())
            self.reset_lazy_loading_state()
            self.view.ui.tableView.scrollToTop()
            if self.admin_comein == True:
                room_id = self.view.get_type_search()
            else:
                room_id = self.log_room_id if hasattr(self, 'log_room_id') else None
            self.load_lab_orders_data(room_id, self.current_offset, self.limit)
        else:
            if len(barcode) == 12:
                barcode = barcode.lstrip('0')
                self.search_by_barcode(barcode)
            else:
                QMessageBox.warning(self.view, "Barcode Error", "เลข Barcode ต้องมีจำนวน 12 ตัวเลขเท่านั้น.")
                self.view.ui.barcode_lineEdit.clear()
                return
    
    def search_by_barcode(self, barcode):
        """Search lab orders by barcode"""
        self.model.removeRows(0, self.model.rowCount())
        self.view.ui.tableView.scrollToTop()
        try:
            room_id_param = "" if self.admin_comein else str(self.log_room_id)
            result = self.receive_lab_service.get_lab_order_by_barcode(barcode, room_id_param)
            if result and result.get('found', False):
                job_progress = result['job_progress']
                self.all_data = job_progress
                for item in job_progress:
                    time_item = QStandardItem(str(item.get('time', '')))
                    time_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    lab_order_id = str(item.get('lab_order_id', ''))
                    lab_order_id_formatted = lab_order_id.zfill(12)
                    lab_order_id_item = QStandardItem(lab_order_id_formatted)
                    lab_order_id_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    speed_item = QStandardItem(str(item.get('speed', '')))
                    speed_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    sample_inspection_item = QStandardItem(str(item.get('sample_inspection', '')))
                    sample_inspection_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    row = [time_item, lab_order_id_item, speed_item, sample_inspection_item]
                    self.model.appendRow(row)
                QMessageBox.information(self.view, "Success", result.get('message', f"พบข้อมูล {len(job_progress)} รายการ"))
            else:
                message = result.get('message', 'ไม่พบข้อมูล Barcode นี้') if result else 'ไม่พบข้อมูล Barcode นี้'
                QMessageBox.warning(self.view, "Not Found", message)
                
        except Exception as e:
            QMessageBox.warning(self.view, "Search Error", f"ไม่สามารถค้นหาข้อมูลได้: {str(e)}")
    
    def load_more_lab_orders(self):
        if self.is_loading or not self.has_more_data:
            return
        room_id = self.log_room_id if hasattr(self, 'log_room_id') else None
        self.load_lab_orders_data(room_id, self.current_offset, self.limit)
    
    def load_lab_orders_data(self, room_id, offset, limit):
        if room_id is None:
            return
        self.is_loading = True
        try:
            result = self.receive_lab_service.get_lab_order_to_day(
                room_id=room_id, 
                offset=offset, 
                limit=limit
            )
            if result and 'job_progress' in result:
                job_progress = result['job_progress']
                self.has_more_data = result.get('has_more', False)
                for item in job_progress:
                    time_item = QStandardItem(str(item.get('time', '')))
                    time_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    lab_order_id = str(item.get('lab_order_id', ''))
                    lab_order_id_formatted = lab_order_id.zfill(12)
                    lab_order_id_item = QStandardItem(lab_order_id_formatted)
                    lab_order_id_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    speed_item = QStandardItem(str(item.get('speed', '')))
                    speed_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    sample_inspection_item = QStandardItem(str(item.get('sample_inspection', '')))
                    sample_inspection_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    row = [time_item, lab_order_id_item, speed_item, sample_inspection_item]
                    self.model.appendRow(row)
                    self.all_data.append(item)
                self.current_offset += len(job_progress)
            else:
                self.has_more_data = False
        except Exception as e:
            QMessageBox.warning(self.view, "Load Error", f"ไม่สามารถโหลดข้อมูลได้: {str(e)}")
            self.has_more_data = False
        finally:
            self.is_loading = False

    # ==========================================
    # ROOM & ACCESS CONTROL - จัดการห้องและการเข้าถึง
    # ==========================================
    def _set_room_for_user(self, room, room_id):
        self.log_room = room
        self.log_room_id = room_id
        if self.log_room == "ห้องปฏิบัติการส่วนกลาง":
            self.admin_comein = True
            self.view.show_radio_buttons()
            self.log_room_id = self.view.get_type_search()
        else:
            self.admin_comein = False
            self.view.hide_radio_buttons()
        return self.log_room, self.log_room_id
    