import base64
import os
import sys
import pythoncom
import shutil
from datetime import datetime
from docx import Document
from PySide6.QtCore import QObject, Qt, QUrl, QThread, Signal
from PySide6.QtWidgets import QMessageBox, QAbstractItemView, QFileDialog, QProgressDialog
from PySide6.QtGui import QStandardItemModel, QStandardItem
from PySide6.QtWebEngineCore import QWebEngineSettings
from docx2pdf import convert
from View.view_monocular_form import MonocularFormView
from SERVICES_REPORT_LAB.send_lab_report_service import SendLabService
from SERVICES_REPORT_LAB.report_information_service import ReportInformationService
from SERVICES_REPORT_LAB.report_form_service import ReportFormService


class ConvertPDFThread(QThread):
    finished_signal = Signal()      
    error_signal = Signal(str)       

    def __init__(self, word_path, pdf_path):
        super().__init__()
        self.word_path = word_path
        self.pdf_path = pdf_path

    def run(self):
        try:
            import sys
            import os
            
            pythoncom.CoInitialize()
            
            # Suppress docx2pdf progress bar output
            # Redirect stderr to devnull to hide progress bar
            with open(os.devnull, 'w') as devnull:
                old_stderr = sys.stderr
                sys.stderr = devnull
                try:
                    convert(self.word_path, self.pdf_path)
                finally:
                    sys.stderr = old_stderr
            
            self.finished_signal.emit()
            
        except Exception as e:
            self.error_signal.emit(str(e))
            
        finally:
            pythoncom.CoUninitialize()


class MonocularFormController(QObject):
    """ Controller for the Monocular Form Page """

    def __init__(self, view: MonocularFormView, main_controller=None):
        super().__init__()
        self.view: MonocularFormView = view
        self.main_controller = main_controller
        self.send_lab_service = SendLabService()
        self.report_info_service = ReportInformationService()
        self.report_form_service = ReportFormService()
        
        # --- State Variables for Barcode List ---
        self.current_offset = 0
        self.limit = 50
        self.is_loading = False
        self.has_more_data = True
        self.all_data = []
        
        # Room ID for Molecular Biology
        self.room_id = 8
        
        # Selection Data
        self.selected_lab_order_id = None
        self.selected_form_type = None
        
        # PDF Conversion State
        self.pdf_output_path = None
        self.convert_thread = None
        
        # Filled Form Path (for saving)
        self.filled_word_path = None
        self.current_report_form_id = None
        
        # Progress Dialog
        self.progress_dialog = None
        
        # --- Temp Folder Setup ---
        if getattr(sys, 'frozen', False):
            base_dir = os.path.dirname(sys.executable)
        else:
            current_script_path = os.path.abspath(__file__)
            controller_dir = os.path.dirname(current_script_path)
            project_dir = os.path.dirname(controller_dir)
            base_dir = os.path.dirname(project_dir)
            
        self.temp_report_dir = os.path.join(base_dir, 'temp_file_report')
        if not os.path.exists(self.temp_report_dir):
            try:
                os.makedirs(self.temp_report_dir)
            except OSError as e:
                fallback_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'temp_file_report')
                os.makedirs(fallback_dir, exist_ok=True)
                self.temp_report_dir = fallback_dir
        
        # Setup UI components
        self.setup_barcode_table_model()
        self._setup_ui()
        self._setup_connections()
        
        # Load initial data
        self.load_barcode_data()
    
    def setup_barcode_table_model(self):
        """ตั้งค่าตารางรายการบาร์โค้ด"""
        self.barcode_model = QStandardItemModel()
        self.barcode_model.setHorizontalHeaderLabels(['Lab Order ID'])
        self.view.ui.barcode_mono_tableView.setModel(self.barcode_model)
        self.view.ui.barcode_mono_tableView.setColumnWidth(0, 200)
        self.view.ui.barcode_mono_tableView.horizontalHeader().setStretchLastSection(True)
        self.view.ui.barcode_mono_tableView.verticalHeader().setVisible(False)
        self.view.ui.barcode_mono_tableView.setShowGrid(True)
        self.view.ui.barcode_mono_tableView.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.view.ui.barcode_mono_tableView.setSelectionMode(QAbstractItemView.SingleSelection)
        self.view.ui.barcode_mono_tableView.setAlternatingRowColors(True)
        self.view.ui.barcode_mono_tableView.setEditTriggers(QAbstractItemView.NoEditTriggers)
    
    def _setup_ui(self):
        """ตั้งค่า UI components"""
        # Setup Barcode LineEdit
        self.view.ui.barcode_mono_lineEdit.setPlaceholderText("พิมพ์บาร์โค้ดเพื่อค้นหา")
        self.view.ui.barcode_mono_lineEdit.clear()
        
        # Set default radio button to PCR FIP
        self.view.ui.PCR_FIP_radioButton.setChecked(True)
        self.selected_form_type = 'PCR FIP'
        
    def _setup_connections(self):
        """เชื่อมต่อ signals และ slots"""
        # Connect Barcode Table and Search
        self.view.ui.barcode_mono_tableView.doubleClicked.connect(self.on_barcode_double_clicked)
        self.view.ui.select_barcode_mono__pushButton.clicked.connect(self.on_select_barcode_clicked)
        self.view.ui.barcode_mono_lineEdit.returnPressed.connect(self.on_barcode_search)
        
        # Connect Form Preview Button
        self.view.ui.form_preview_pushButton.clicked.connect(self.on_form_preview_clicked)
        
        # Connect Save Button
        self.view.ui.save_pushButton.clicked.connect(self.on_save_button_clicked)
        
        # Connect Radio Buttons
        self.view.ui.PCR_FIP_radioButton.toggled.connect(lambda checked: self.on_radio_changed('PCR FIP', checked))
        self.view.ui.PCR_real_time_radioButton.toggled.connect(lambda checked: self.on_radio_changed('PCR real time AHS', checked))
        self.view.ui.PCR_FeLV_radioButton.toggled.connect(lambda checked: self.on_radio_changed('PCR FeLV', checked))
        self.view.ui.PCR_ASF_radioButton.toggled.connect(lambda checked: self.on_radio_changed('PCR real time ASF', checked))
        
        # Connect scroll for infinite loading
        scrollbar = self.view.ui.barcode_mono_tableView.verticalScrollBar()
        scrollbar.valueChanged.connect(self.on_scroll)
    
    def on_radio_changed(self, form_type, checked):
        """Event handler เมื่อเลือก Radio Button"""
        if checked:
            self.selected_form_type = form_type
    
    def on_form_preview_clicked(self):
        """Event handler เมื่อกดปุ่ม Form Preview"""
        if not self.selected_form_type:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณาเลือกประเภทฟอร์มก่อน")
            return
        
        if not self.selected_lab_order_id:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณาเลือกบาร์โค้ดก่อน")
            return
        
        # สร้าง Progress Dialog
        self.progress_dialog = QProgressDialog("กำลังโหลดฟอร์ม...", None, 0, 100, self.view)
        self.progress_dialog.setWindowTitle("กำลังประมวลผล")
        self.progress_dialog.setWindowModality(Qt.WindowModal)
        self.progress_dialog.setMinimumDuration(0)
        self.progress_dialog.setValue(0)
        self.progress_dialog.show()
        
        # Mapping form type to report name
        form_name_mapping = {
            'PCR FIP': 'รายงานผล PCR_FIP.docx',
            'PCR real time AHS': 'รายงานผล PCR_realtime_AHS.docx',
            'PCR FeLV': 'รายงานผล PCR_FeLV.docx',
            'PCR real time ASF': 'รายงานผล PCR_realtime_ASF.docx'
        }
        
        report_name = form_name_mapping.get(self.selected_form_type)
        
        if not report_name:
            QMessageBox.warning(self.view, "ข้อผิดพลาด", "ไม่พบชื่อฟอร์มที่ตรงกัน")
            return
        
        # Get report data from database
        try:
            self.progress_dialog.setLabelText("กำลังดึงข้อมูลฟอร์ม...")
            self.progress_dialog.setValue(10)
            reports = self.report_info_service.get_reports_by_room_and_status(self.room_id, status=1)
            
            if not reports:
                QMessageBox.warning(
                    self.view, 
                    "ข้อผิดพลาด", 
                    f"ไม่พบข้อมูลฟอร์มในระบบ\nRoom ID: {self.room_id}"
                )
                return
            
            # Find the matching report
            selected_report = None
            for report in reports:
                current_name = report.get('report_name', '').strip()
                target_name = report_name.strip()
                if current_name == target_name:
                    selected_report = report
                    break
            
            if not selected_report:
                available_names = "\n".join([f"- {r.get('report_name')}" for r in reports])
                QMessageBox.warning(
                    self.view, 
                    "ไม่พบฟอร์ม", 
                    f"ไม่พบฟอร์ม: {report_name}\n\nฟอร์มที่มีในระบบ:\n{available_names}"
                )
                return
            
            # Get report path
            report_path = selected_report.get('report_path')
            
            if not report_path or not os.path.exists(report_path):
                if self.progress_dialog:
                    self.progress_dialog.close()
                QMessageBox.warning(
                    self.view, 
                    "ข้อผิดพลาด", 
                    f"ไม่พบไฟล์ฟอร์ม:\n{report_path}"
                )
                return
            
            # Fill data in Word document
            self.progress_dialog.setLabelText("กำลังเติมข้อมูลในฟอร์ม...")
            self.progress_dialog.setValue(30)
            filled_path = self.fill_form_data(report_path, report_name)
            
            if not filled_path:
                if self.progress_dialog:
                    self.progress_dialog.close()
                QMessageBox.warning(
                    self.view,
                    "ข้อผิดพลาด",
                    "ไม่สามารถเติมข้อมูลในฟอร์มได้"
                )
                return
            
            # Save filled path for later use
            self.filled_word_path = filled_path
            
            # Convert Word to PDF
            self.progress_dialog.setLabelText("กำลังแปลงเป็น PDF...")
            self.progress_dialog.setValue(70)
            self.convert_form_to_pdf(filled_path, report_name)
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            if self.progress_dialog:
                self.progress_dialog.close()
            QMessageBox.critical(
                self.view, 
                "ข้อผิดพลาด", 
                f"เกิดข้อผิดพลาดในการโหลดฟอร์ม:\n{str(e)}"
            )
    
    def fill_form_data(self, template_path, report_name):
        """เติมข้อมูลลงในฟอร์ม Word"""
        try:
            # Get lab order details
            lab_order_id = int(self.selected_lab_order_id)
            lab_order_data = self.report_form_service.get_lab_order_details(lab_order_id)
            
            if not lab_order_data:
                return None
            
            # NOTE: ไม่ต้องดึง test_items เพราะไม่เติมตาราง - เติมเฉพาะส่วนหัว
            # lab_order_with_tests = self.report_form_service.get_lab_order_with_tests(lab_order_id)
            # test_items = []
            # if lab_order_with_tests and lab_order_with_tests.get('success'):
            #     test_items = lab_order_with_tests.get('test_items', [])
            
            # Get next report form ID
            report_id_data = self.report_form_service.get_latest_id()
            next_report_id = report_id_data.get('next_id', 1)
            
            # Generate sample number (เลขที่ตัวอย่าง)
            # Format: D-{day} {lab_order_id} เช่น "D-23 1605"
            current_date = datetime.now()
            day = current_date.day
            sample_number = f"D-{day} {lab_order_id}"
            
            # Helper function to format date in Thai Buddhist calendar
            def format_thai_date(dt):
                """Format datetime to Thai Buddhist calendar: d/m/YYYY"""
                if isinstance(dt, str):
                    try:
                        # Try ISO format with T (e.g., 2025-12-08T12:40:06)
                        dt = datetime.fromisoformat(dt.replace('Z', '+00:00'))
                    except:
                        try:
                            dt = datetime.strptime(dt, "%Y-%m-%d %H:%M:%S")
                        except:
                            try:
                                dt = datetime.strptime(dt, "%Y-%m-%d")
                            except:
                                dt = datetime.now()
                # Thai Buddhist year = Gregorian year + 543
                thai_year = dt.year + 543
                # Format as d/m/yyyy (without leading zeros for day)
                return f"{dt.day}/{dt.month}/{thai_year}"
            
            # วันที่รับตัวอย่าง - มาจาก case_registration.dtime
            case_dtime = lab_order_data.get('case_dtime')
            if case_dtime:
                receive_date = format_thai_date(case_dtime)
            else:
                receive_date = format_thai_date(current_date)
            
            # วันที่เริ่มทดสอบ - มาจาก lab_receive_detail.dtime (วันที่รับเคส)
            receive_dtime = lab_order_data.get('receive_dtime')
            if receive_dtime:
                test_start_date = format_thai_date(receive_dtime)
            else:
                test_start_date = ""  # ไม่มีข้อมูลให้เว้นว่าง
            
            # เลขที่รายงาน
            report_number = str(next_report_id)
            
            # Load Word document
            doc = Document(template_path)
            
            # วันที่ออบผล (วันที่ปัจจุบัน)
            result_date = format_thai_date(current_date)
            
            # Define replacement mapping
            replacements = {
                '{{เลขที่รายงาน}}': report_number,
                '{{เลขที่ตัวอย่าง}}': sample_number,
                '{{วันที่รับตัวอย่าง}}': receive_date,
                '{{วันที่เริ่มทดสอบ}}': test_start_date,
                '{{วันที่เริ่มตรวจคลอง}}': test_start_date,  # เหมือนกับวันที่เริ่มทดสอบ
                '{{วันที่ออบผล}}': result_date,
            }
            
            # Helper function to replace text in runs (preserves formatting)
            def replace_in_paragraph(paragraph, replacements):
                """Replace placeholders in paragraph while preserving formatting"""
                from docx.shared import Pt
                
                for key, value in replacements.items():
                    if key in paragraph.text:
                        full_text = paragraph.text
                        
                        if key in full_text:
                            new_text = full_text.replace(key, value)
                            
                            template_run = None
                            for run in paragraph.runs:
                                if run.text.strip():
                                    template_run = run
                                    break
                            
                            for run in paragraph.runs:
                                run.text = ''
                            
                            if paragraph.runs:
                                target_run = paragraph.runs[0]
                            else:
                                target_run = paragraph.add_run()
                            
                            target_run.text = new_text
                            target_run.font.name = 'TH SarabunPSK'
                            target_run.font.size = Pt(16)
                            
                            if template_run is not None:
                                target_run.font.bold = template_run.font.bold
                                target_run.font.italic = template_run.font.italic
                                target_run.font.underline = template_run.font.underline
                                if template_run.font.color.rgb:
                                    target_run.font.color.rgb = template_run.font.color.rgb
            
            # NOTE: ไม่ต้องเติมตารางรายการทดสอบ - เติมเฉพาะส่วนหัวเท่านั้น
            # (ตามคำขอของผู้ใช้: เติมแค่ส่วนหัว ไม่ต้องเติมตาราง)
            
            # Helper function to replace in all tables and paragraphs
            def replace_in_all_tables(element, replacements):
                """Recursively replace in all tables including nested ones"""
                if hasattr(element, 'tables'):
                    for table in element.tables:
                        # ไม่ต้องเติม test items ในตาราง - เติมเฉพาะ placeholders ในส่วนหัว
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    replace_in_paragraph(paragraph, replacements)
                                replace_in_all_tables(cell, replacements)
            
            # Replace in main paragraphs
            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph, replacements)
            
            # Replace in all tables (including nested) - เติมเฉพาะส่วนหัว
            replace_in_all_tables(doc, replacements)
            
            # Save filled document to temp folder
            filled_filename = f"filled_{report_name}"
            filled_path = os.path.join(self.temp_report_dir, filled_filename)
            
            # Remove old file if exists
            if os.path.exists(filled_path):
                os.remove(filled_path)
            
            doc.save(filled_path)
            
            # Get recorder (group_id) from logged-in user
            recorder_id = 0
            if self.main_controller and hasattr(self.main_controller, 'logged_in_user_info'):
                user_info = self.main_controller.logged_in_user_info
                if user_info and isinstance(user_info, dict):
                    recorder_id = user_info.get('group_id', 0)
            
            # Save to database
            success, report_form_id, message = self.report_form_service.create_report_form(
                lab_order_id=lab_order_id,
                location="",
                comment="",
                state=1,
                status=1,
                recorder=recorder_id,
                approver=0,
                room_id=8  # Molecular Biology (PCR)
            )
            
            if success:
                self.current_report_form_id = report_form_id
            
            return filled_path
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return None
    
    def on_save_button_clicked(self):
        """Event handler เมื่อกดปุ่มบันทึก - บันทึกไฟล์ .docx"""
        if not self.filled_word_path or not os.path.exists(self.filled_word_path):
            QMessageBox.warning(
                self.view,
                "แจ้งเตือน",
                "กรุณากดปุ่ม FORM PREVIEW ก่อนบันทึก"
            )
            return
        
        try:
            default_filename = f"Monocular_Report_{self.selected_lab_order_id}.docx"
            
            file_path, _ = QFileDialog.getSaveFileName(
                self.view,
                "บันทึกไฟล์ฟอร์ม",
                default_filename,
                "Word Documents (*.docx)"
            )
            
            if file_path:
                if not file_path.lower().endswith('.docx'):
                    file_path += '.docx'
                
                shutil.copy2(self.filled_word_path, file_path)
                
                if self.current_report_form_id:
                    self.report_form_service.update_report_form(
                        report_id=self.current_report_form_id,
                        location=file_path
                    )
                
                QMessageBox.information(
                    self.view,
                    "สำเร็จ",
                    f"บันทึกไฟล์สำเร็จ\n{file_path}"
                )
        
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(
                self.view,
                "ข้อผิดพลาด",
                f"ไม่สามารถบันทึกไฟล์ได้:\n{str(e)}"
            )
    
    def load_barcode_data(self):
        """โหลดข้อมูลบาร์โค้ดจาก API"""
        self.is_loading = True
        try:
            result = self.send_lab_service.get_received_labs_to_day(str(self.room_id), self.current_offset, self.limit)
            if result and 'job_progress' in result:
                self.has_more_data = result.get('has_more', False)
                for item in result['job_progress']:
                    self._add_row_to_table(item)
                    self.all_data.append(item)
                self.current_offset += len(result['job_progress'])
            else:
                self.has_more_data = False
        except Exception as e:
            QMessageBox.warning(self.view, "ข้อผิดพลาด", f"ไม่สามารถโหลดข้อมูลบาร์โค้ดได้: {str(e)}")
        finally:
            self.is_loading = False
    
    def _add_row_to_table(self, item):
        """เพิ่มข้อมูลลงในตาราง"""
        lab_id = str(item.get('lab_order_id', '')).zfill(12)
        item_obj = QStandardItem(lab_id)
        item_obj.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        self.barcode_model.appendRow([item_obj])
    
    def on_scroll(self, value):
        """Event handler สำหรับ infinite scroll"""
        scrollbar = self.view.ui.barcode_mono_tableView.verticalScrollBar()
        if value >= scrollbar.maximum() - 10:
            if self.has_more_data and not self.is_loading:
                if not self.view.ui.barcode_mono_lineEdit.text():
                    self.load_more_data()
    
    def load_more_data(self):
        """โหลดข้อมูลเพิ่มเติม"""
        self.load_barcode_data()
    
    def on_barcode_double_clicked(self, index):
        """Event handler เมื่อ double click ที่ตาราง"""
        if index.isValid():
            row = index.row()
            barcode = self.barcode_model.item(row, 0).text()
            self.view.ui.barcode_mono_lineEdit.setText(barcode)
            self.selected_lab_order_id = barcode
    
    def on_select_barcode_clicked(self):
        """Event handler เมื่อกดปุ่มเลือกบาร์โค้ด"""
        barcode = self.view.ui.barcode_mono_lineEdit.text().strip()
        if barcode:
            self.selected_lab_order_id = barcode
        else:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณากรอกหรือเลือกบาร์โค้ด")
    
    def on_barcode_search(self):
        """Event handler เมื่อกด Enter ในช่องค้นหา"""
        barcode = self.view.ui.barcode_mono_lineEdit.text().strip()
        if barcode:
            self.search_by_barcode(barcode)
    
    def search_by_barcode(self, barcode):
        """ค้นหาบาร์โค้ด"""
        self.barcode_model.removeRows(0, self.barcode_model.rowCount())
        self.view.ui.barcode_mono_tableView.scrollToTop()
        
        try:
            result = self.send_lab_service.get_received_labs_by_barcode(barcode, str(self.room_id))
            
            if result and result.get('found', False):
                for item in result['job_progress']:
                    self._add_row_to_table(item)
            else:
                QMessageBox.warning(self.view, "ไม่พบข้อมูล", "ไม่พบบาร์โค้ดนี้ในระบบ")
        except Exception as e:
            QMessageBox.warning(self.view, "ข้อผิดพลาด", f"เกิดข้อผิดพลาดในการค้นหา: {str(e)}")
    
    def reset_lazy_loading_state(self):
        """รีเซ็ตสถานะการโหลดข้อมูล"""
        self.current_offset = 0
        self.is_loading = False
        self.has_more_data = True
        self.all_data = []
    
    def reload_data(self):
        """โหลดข้อมูลใหม่"""
        self.barcode_model.removeRows(0, self.barcode_model.rowCount())
        self.reset_lazy_loading_state()
        self.view.ui.barcode_mono_lineEdit.clear()
        self.load_barcode_data()
    
    def get_selected_barcode(self):
        """ดึงค่าบาร์โค้ดที่เลือก"""
        return self.selected_lab_order_id
    
    def clear_form(self):
        """ล้างค่าในฟอร์ม"""
        self.view.ui.barcode_mono_lineEdit.clear()
        self.selected_lab_order_id = None
    
    def convert_form_to_pdf(self, word_path, report_name):
        """แปลงฟอร์ม Word เป็น PDF และแสดงผล"""
        try:
            pdf_filename = f"preview_{report_name.replace('.docx', '.pdf')}"
            self.pdf_output_path = os.path.join(self.temp_report_dir, pdf_filename)
            
            if os.path.exists(self.pdf_output_path):
                os.remove(self.pdf_output_path)
            
            self.convert_thread = ConvertPDFThread(word_path, self.pdf_output_path)
            self.convert_thread.finished_signal.connect(self.on_conversion_finished)
            self.convert_thread.error_signal.connect(self.on_conversion_error)
            self.convert_thread.start()
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(
                self.view, 
                "ข้อผิดพลาด", 
                f"ไม่สามารถเริ่มแปลงไฟล์ได้:\n{str(e)}"
            )
    
    def on_conversion_finished(self):
        """ทำงานเมื่อแปลงไฟล์สำเร็จ"""
        if self.progress_dialog:
            self.progress_dialog.setValue(100)
            self.progress_dialog.setLabelText("เสร็จสิ้น!")
            self.progress_dialog.close()
        self.display_pdf_preview()
        self.convert_thread = None
    
    def on_conversion_error(self, error_msg):
        """ทำงานเมื่อแปลงไฟล์ Error"""
        QMessageBox.critical(
            self.view, 
            "ข้อผิดพลาด", 
            f"ไม่สามารถแปลงไฟล์ได้:\n{error_msg}"
        )
        self.convert_thread = None
    
    def display_pdf_preview(self):
        """แสดงตัวอย่าง PDF ใน WebEngineView"""
        if self.pdf_output_path and os.path.exists(self.pdf_output_path):
            try:
                abs_pdf_path = os.path.abspath(self.pdf_output_path)
                
                settings = self.view.ui.preview_word_webEngineView.page().settings()
                settings.setAttribute(QWebEngineSettings.WebAttribute.PluginsEnabled, True)
                settings.setAttribute(QWebEngineSettings.WebAttribute.PdfViewerEnabled, True)
                
                with open(abs_pdf_path, "rb") as f:
                    pdf_data = f.read()
                    base64_pdf = base64.b64encode(pdf_data).decode('utf-8')
                
                pdf_filename = os.path.basename(abs_pdf_path)
                html_content = self.create_pdf_viewer_html(base64_pdf, pdf_filename)
                self.view.ui.preview_word_webEngineView.setHtml(html_content, QUrl("http://localhost"))
                
            except Exception as e:
                import traceback
                traceback.print_exc()
                QMessageBox.warning(
                    self.view, 
                    "ข้อผิดพลาด", 
                    f"ไม่สามารถแสดงตัวอย่าง PDF ได้:\n{str(e)}"
                )
        else:
            QMessageBox.warning(
                self.view,
                "ข้อผิดพลาด",
                f"ไม่พบไฟล์ PDF ที่แปลงแล้ว\n{self.pdf_output_path}"
            )
    
    def create_pdf_viewer_html(self, base64_data, filename="document.pdf"):
        """สร้าง HTML สำหรับแสดง PDF"""
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{filename}</title>
            <style>
                body, html {{ margin: 0; padding: 0; height: 100%; overflow: hidden; background-color: #525659; }}
                embed {{ width: 100%; height: 100%; border: none; }}
            </style>
        </head>
        <body>
            <embed 
                src="data:application/pdf;base64,{base64_data}#toolbar=1&navpanes=0&scrollbar=1&page=1&view=FitH" 
                type="application/pdf" width="100%" height="100%"
            >
        </body>
        </html>
        """
