import base64
import os
import sys
import pythoncom
import shutil
from datetime import datetime
from docx import Document
from PySide6.QtCore import QObject, Qt, QUrl, QThread, Signal
from PySide6.QtWidgets import QMessageBox, QAbstractItemView, QFileDialog, QProgressDialog
from PySide6.QtGui import QStandardItemModel, QStandardItem
from PySide6.QtWebEngineCore import QWebEngineSettings
from docx2pdf import convert
from View.view_parasite_form import ParasiteFormView
from SERVICES_REPORT_LAB.send_lab_report_service import SendLabService
from SERVICES_REPORT_LAB.report_information_service import ReportInformationService
from SERVICES_REPORT_LAB.report_form_service import ReportFormService


class ConvertPDFThread(QThread):
    finished_signal = Signal()      
    error_signal = Signal(str)       

    def __init__(self, word_path, pdf_path):
        super().__init__()
        self.word_path = word_path
        self.pdf_path = pdf_path

    def run(self):
        try:
            import sys
            import os
            
            pythoncom.CoInitialize()
            
            # Suppress docx2pdf progress bar output
            # Redirect stderr to devnull to hide progress bar
            with open(os.devnull, 'w') as devnull:
                old_stderr = sys.stderr
                sys.stderr = devnull
                try:
                    convert(self.word_path, self.pdf_path)
                finally:
                    sys.stderr = old_stderr
            
            self.finished_signal.emit()
            
        except Exception as e:
            self.error_signal.emit(str(e))
            
        finally:
            pythoncom.CoUninitialize()


class ParasiteFormController(QObject):
    """ Controller for the Parasite Form Page """

    def __init__(self, view: ParasiteFormView, main_controller=None):
        super().__init__()
        self.view: ParasiteFormView = view
        self.main_controller = main_controller
        self.send_lab_service = SendLabService()
        self.report_info_service = ReportInformationService()
        self.report_form_service = ReportFormService()
        
        # --- State Variables for Barcode List ---
        self.current_offset = 0
        self.limit = 50
        self.is_loading = False
        self.has_more_data = True
        self.all_data = []
        
        # Room ID for Parasitology
        self.room_id = 5
        
        # Selection Data
        self.selected_lab_order_id = None
        self.selected_form_type = None
        
        # PDF Conversion State
        self.pdf_output_path = None
        self.convert_thread = None
        
        # Filled Form Path (for saving)
        self.filled_word_path = None
        self.current_report_form_id = None
        
        # Progress Dialog
        self.progress_dialog = None
        
        # --- Temp Folder Setup ---
        if getattr(sys, 'frozen', False):
            base_dir = os.path.dirname(sys.executable)
        else:
            current_script_path = os.path.abspath(__file__)
            controller_dir = os.path.dirname(current_script_path)
            project_dir = os.path.dirname(controller_dir)
            base_dir = os.path.dirname(project_dir)
            
        self.temp_report_dir = os.path.join(base_dir, 'temp_file_report')
        if not os.path.exists(self.temp_report_dir):
            try:
                os.makedirs(self.temp_report_dir)
                # DEBUG: Uncomment for debugging
                # print(f"Created Temp Dir at: {self.temp_report_dir}")
            except OSError as e:
                # DEBUG: Uncomment for debugging
                # print(f"Error creating temp dir: {e}")
                fallback_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'temp_file_report')
                os.makedirs(fallback_dir, exist_ok=True)
                self.temp_report_dir = fallback_dir
        
        # Setup UI components
        self.setup_barcode_table_model()
        self._setup_ui()
        self._setup_connections()
        
        # Load initial data
        self.load_barcode_data()
    
    def setup_barcode_table_model(self):
        """ตั้งค่าตารางรายการบาร์โค้ด"""
        self.barcode_model = QStandardItemModel()
        self.barcode_model.setHorizontalHeaderLabels(['Lab Order ID'])
        self.view.ui.barcode_para_tableView.setModel(self.barcode_model)
        self.view.ui.barcode_para_tableView.setColumnWidth(0, 200)
        self.view.ui.barcode_para_tableView.horizontalHeader().setStretchLastSection(True)
        self.view.ui.barcode_para_tableView.verticalHeader().setVisible(False)
        self.view.ui.barcode_para_tableView.setShowGrid(True)
        self.view.ui.barcode_para_tableView.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.view.ui.barcode_para_tableView.setSelectionMode(QAbstractItemView.SingleSelection)
        self.view.ui.barcode_para_tableView.setAlternatingRowColors(True)
        self.view.ui.barcode_para_tableView.setEditTriggers(QAbstractItemView.NoEditTriggers)
    
    def _setup_ui(self):
        """ตั้งค่า UI components"""
        # Setup Color ComboBox (comboBox - ตัวแรกซ้ายสุด)
        self.setup_color_combobox()
        
        # Setup Consistency ComboBox (comboBox_2 - ตัวที่สองตรงกลาง)
        self.setup_consistency_combobox()
        
        # Setup Method ComboBox (comboBox_3 - ตัวที่สามขวาสุด)
        self.setup_method_combobox()
        
        # Setup Barcode LineEdit
        self.view.ui.barcode_para_lineEdit.setPlaceholderText("พิมพ์บาร์โค้ดเพื่อค้นหา")
        self.view.ui.barcode_para_lineEdit.clear()
        
        # Set default radio button to Faces
        self.view.ui.Faces_radioButton.setChecked(True)
        self.selected_form_type = 'Faces'
        
        # Enable all dropdowns for Faces (default)
        self.update_dropdown_states('Faces')
        
    def setup_color_combobox(self):
        """ตั้งค่า dropdown สำหรับ Color"""
        color_options = [
            "Black",
            "Brown",
            "Yellow",
            "Green",
            "Others"
        ]
        
        # Clear existing items
        self.view.ui.comboBox.clear()
        
        # Add items to combobox
        self.view.ui.comboBox.addItems(color_options)
        
        # Set placeholder or default
        self.view.ui.comboBox.setCurrentIndex(-1)  # No selection by default
        self.view.ui.comboBox.setPlaceholderText("Select Color")
    
    def setup_consistency_combobox(self):
        """ตั้งค่า dropdown สำหรับ Consistency"""
        consistency_options = [
            "Hard",
            "Formed",
            "Soft",
            "Mushy",
            "Loose",
            "Diarrhoeic",
            "Watery"
        ]
        
        # Clear existing items
        self.view.ui.comboBox_2.clear()
        
        # Add items to combobox
        self.view.ui.comboBox_2.addItems(consistency_options)
        
        # Set placeholder or default
        self.view.ui.comboBox_2.setCurrentIndex(-1)  # No selection by default
        self.view.ui.comboBox_2.setPlaceholderText("Select Consistency")
    
    def setup_method_combobox(self):
        """ตั้งค่า dropdown สำหรับ Method"""
        method_options = [
            "Flotation",
            "Sedimentation",
            "Others"
        ]
        
        # Clear existing items
        self.view.ui.comboBox_3.clear()
        
        # Add items to combobox
        self.view.ui.comboBox_3.addItems(method_options)
        
        # Set placeholder or default
        self.view.ui.comboBox_3.setCurrentIndex(-1)  # No selection by default
        self.view.ui.comboBox_3.setPlaceholderText("Select Method")
    
    def _setup_connections(self):
        """เชื่อมต่อ signals และ slots"""
        # Connect Color ComboBox
        self.view.ui.comboBox.currentTextChanged.connect(self.on_color_changed)
        
        # Connect Consistency ComboBox
        self.view.ui.comboBox_2.currentTextChanged.connect(self.on_consistency_changed)
        
        # Connect Method ComboBox
        self.view.ui.comboBox_3.currentTextChanged.connect(self.on_method_changed)
        
        # Connect Barcode Table and Search
        self.view.ui.barcode_para_tableView.doubleClicked.connect(self.on_barcode_double_clicked)
        self.view.ui.select_barcode_para__pushButton.clicked.connect(self.on_select_barcode_clicked)
        self.view.ui.barcode_para_lineEdit.returnPressed.connect(self.on_barcode_search)
        
        # Connect Form Preview Button
        self.view.ui.form_preview_pushButton.clicked.connect(self.on_form_preview_clicked)
        
        # Connect Save Button
        self.view.ui.save_pushButton.clicked.connect(self.on_save_button_clicked)
        
        # Connect Radio Buttons
        self.view.ui.Faces_radioButton.toggled.connect(lambda checked: self.on_radio_changed('Faces', checked))
        self.view.ui.Faces_dog_cat_radioButton.toggled.connect(lambda checked: self.on_radio_changed('Faces Dog Cat', checked))
        self.view.ui.radioButton_3.toggled.connect(lambda checked: self.on_radio_changed('Blood', checked))
        self.view.ui.radioButton_4.toggled.connect(lambda checked: self.on_radio_changed('Identification', checked))
        
        # Connect scroll for infinite loading
        scrollbar = self.view.ui.barcode_para_tableView.verticalScrollBar()
        scrollbar.valueChanged.connect(self.on_scroll)
    
    def on_color_changed(self, text):
        """Event handler เมื่อเลือก Color"""
        if text:
            print(f"Selected Color: {text}")
    
    def on_consistency_changed(self, text):
        """Event handler เมื่อเลือก Consistency"""
        if text:
            print(f"Selected Consistency: {text}")
    
    def on_method_changed(self, text):
        """Event handler เมื่อเลือก Method"""
        if text:
            print(f"Selected Method: {text}")
    
    def on_radio_changed(self, form_type, checked):
        """Event handler เมื่อเลือก Radio Button"""
        if checked:
            self.selected_form_type = form_type
            # Update dropdown states based on form type
            self.update_dropdown_states(form_type)
    
    def update_dropdown_states(self, form_type):
        """Enable/Disable dropdowns ตามประเภทฟอร์ม"""
        # Faces และ Faces Dog Cat → Enable ทั้งหมด
        # Blood และ Identification → Disable ทั้งหมด
        
        if form_type in ['Faces', 'Faces Dog Cat']:
            # Enable all dropdowns
            self.view.ui.comboBox.setEnabled(True)
            self.view.ui.comboBox_2.setEnabled(True)
            self.view.ui.comboBox_3.setEnabled(True)
        else:  # Blood or Identification
            # Disable all dropdowns
            self.view.ui.comboBox.setEnabled(False)
            self.view.ui.comboBox_2.setEnabled(False)
            self.view.ui.comboBox_3.setEnabled(False)
            # Clear selections
            self.view.ui.comboBox.setCurrentIndex(-1)
            self.view.ui.comboBox_2.setCurrentIndex(-1)
            self.view.ui.comboBox_3.setCurrentIndex(-1)
    
    def on_form_preview_clicked(self):
        """Event handler เมื่อกดปุ่ม Form Preview"""
        if not self.selected_form_type:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณาเลือกประเภทฟอร์มก่อน")
            return
        
        if not self.selected_lab_order_id:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณาเลือกบาร์โค้ดก่อน")
            return
        
        # สร้าง Progress Dialog
        self.progress_dialog = QProgressDialog("กำลังโหลดฟอร์ม...", None, 0, 100, self.view)
        self.progress_dialog.setWindowTitle("กำลังประมวลผล")
        self.progress_dialog.setWindowModality(Qt.WindowModal)
        self.progress_dialog.setMinimumDuration(0)
        self.progress_dialog.setValue(0)
        self.progress_dialog.show()
        
        # Mapping form type to report name
        form_name_mapping = {
            'Faces': 'รายงานผล Parasite_feces.docx',
            'Faces Dog Cat': 'รายงานผล Parasite_feces dog_cat.docx',
            'Blood': 'รายงานผล Parasite_blood.docx',
            'Identification': 'รายงานผล Parasite_iden.docx'
        }
        
        report_name = form_name_mapping.get(self.selected_form_type)
        
        if not report_name:
            QMessageBox.warning(self.view, "ข้อผิดพลาด", "ไม่พบชื่อฟอร์มที่ตรงกัน")
            return
        
        # Get report data from database
        try:
            self.progress_dialog.setLabelText("กำลังดึงข้อมูลฟอร์ม...")
            self.progress_dialog.setValue(10)
            reports = self.report_info_service.get_reports_by_room_and_status(self.room_id, status=1)
            
            if not reports:
                QMessageBox.warning(
                    self.view, 
                    "ข้อผิดพลาด", 
                    f"ไม่พบข้อมูลฟอร์มในระบบ\nRoom ID: {self.room_id}"
                )
                return
            
            # Find the matching report
            selected_report = None
            for report in reports:
                current_name = report.get('report_name', '').strip()
                target_name = report_name.strip()
                if current_name == target_name:
                    selected_report = report
                    break
            
            if not selected_report:
                available_names = "\n".join([f"- {r.get('report_name')}" for r in reports])
                QMessageBox.warning(
                    self.view, 
                    "ไม่พบฟอร์ม", 
                    f"ไม่พบฟอร์ม: {report_name}\n\nฟอร์มที่มีในระบบ:\n{available_names}"
                )
                return
            
            # Get report path
            report_path = selected_report.get('report_path')
            
            if not report_path or not os.path.exists(report_path):
                if self.progress_dialog:
                    self.progress_dialog.close()
                QMessageBox.warning(
                    self.view, 
                    "ข้อผิดพลาด", 
                    f"ไม่พบไฟล์ฟอร์ม:\n{report_path}"
                )
                return
            
            # Fill data in Word document
            self.progress_dialog.setLabelText("กำลังเติมข้อมูลในฟอร์ม...")
            self.progress_dialog.setValue(30)
            filled_path = self.fill_form_data(report_path, report_name)
            
            if not filled_path:
                if self.progress_dialog:
                    self.progress_dialog.close()
                QMessageBox.warning(
                    self.view,
                    "ข้อผิดพลาด",
                    "ไม่สามารถเติมข้อมูลในฟอร์มได้"
                )
                return
            
            # Save filled path for later use
            self.filled_word_path = filled_path
            
            # Convert Word to PDF
            self.progress_dialog.setLabelText("กำลังแปลงเป็น PDF...")
            self.progress_dialog.setValue(70)
            self.convert_form_to_pdf(filled_path, report_name)
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            if self.progress_dialog:
                self.progress_dialog.close()
            # DEBUG: Uncomment for debugging
            # print(f"ERROR in on_form_preview_clicked:")
            # print(error_details)
            QMessageBox.critical(
                self.view, 
                "ข้อผิดพลาด", 
                f"เกิดข้อผิดพลาดในการโหลดฟอร์ม:\n{str(e)}"
            )
    
    def fill_form_data(self, template_path, report_name):
        """เติมข้อมูลลงในฟอร์ม Word"""
        try:
            # Get lab order details
            lab_order_id = int(self.selected_lab_order_id)
            lab_order_data = self.report_form_service.get_lab_order_details(lab_order_id)
            
            if not lab_order_data:
                # DEBUG: Uncomment for debugging
                # print(f"ERROR: Cannot get lab order details for ID: {lab_order_id}")
                return None
            
            # Get next report form ID
            report_id_data = self.report_form_service.get_latest_id()
            next_report_id = report_id_data.get('next_id', 1)
            
            # Generate sample number (เลขที่ตัวอย่าง)
            # Format: D-{day} {lab_order_id} เช่น "D-23 1605"
            current_date = datetime.now()
            day = current_date.day
            sample_number = f"D-{day} {lab_order_id}"
            
            # Helper function to format date in Thai Buddhist calendar
            def format_thai_date(dt):
                """Format datetime to Thai Buddhist calendar: d/m/YYYY"""
                if isinstance(dt, str):
                    try:
                        # Try ISO format with T (e.g., 2025-12-08T12:40:06)
                        dt = datetime.fromisoformat(dt.replace('Z', '+00:00'))
                    except:
                        try:
                            dt = datetime.strptime(dt, "%Y-%m-%d %H:%M:%S")
                        except:
                            try:
                                dt = datetime.strptime(dt, "%Y-%m-%d")
                            except:
                                dt = datetime.now()
                # Thai Buddhist year = Gregorian year + 543
                thai_year = dt.year + 543
                # Format as d/m/yyyy (without leading zeros for day)
                return f"{dt.day}/{dt.month}/{thai_year}"
            
            # วันที่รับตัวอย่าง - มาจาก case_registration.dtime
            case_dtime = lab_order_data.get('case_dtime')
            if case_dtime:
                receive_date = format_thai_date(case_dtime)
            else:
                receive_date = format_thai_date(current_date)
            
            # วันที่เริ่มทดสอบ - มาจาก lab_receive_detail.dtime
            receive_dtime = lab_order_data.get('receive_dtime')
            if receive_dtime:
                test_start_date = format_thai_date(receive_dtime)
            else:
                test_start_date = ""  # ไม่มีข้อมูลให้เว้นว่าง
            
            # เลขที่รายงาน
            report_number = str(next_report_id)
            
            # Get selected dropdown values
            selected_color = self.get_selected_color() if self.selected_form_type in ['Faces', 'Faces Dog Cat'] else ""
            selected_consistency = self.get_selected_consistency() if self.selected_form_type in ['Faces', 'Faces Dog Cat'] else ""
            selected_method = self.get_selected_method() if self.selected_form_type in ['Faces', 'Faces Dog Cat'] else ""
            
            # Load Word document
            doc = Document(template_path)
            
            # Define replacement mapping
            replacements = {
                '{{เลขที่รายงาน}}': report_number,
                '{{เลขที่ตัวอย่าง}}': sample_number,
                '{{วันที่รับตัวอย่าง}}': receive_date,
                '{{วันที่เริ่มทดสอบ}}': test_start_date,
                '{{Color}}': selected_color,
                '{{Consistency}}': selected_consistency,
                '{{Method}}': selected_method,
            }
            
            # Helper function to replace text in runs (preserves formatting)
            def replace_in_paragraph(paragraph, replacements):
                """Replace placeholders in paragraph while preserving formatting"""
                from docx.shared import Pt
                
                for key, value in replacements.items():
                    if key in paragraph.text:
                        # Get the full text
                        full_text = paragraph.text
                        
                        if key in full_text:
                            # Replace in the full text
                            new_text = full_text.replace(key, value)
                            
                            # Find a run that has formatting to preserve
                            template_run = None
                            for run in paragraph.runs:
                                if run.text.strip():  # Find first non-empty run
                                    template_run = run
                                    break
                            
                            # Clear all runs
                            for run in paragraph.runs:
                                run.text = ''
                            
                            # Add replaced text to first run (or create new one)
                            if paragraph.runs:
                                target_run = paragraph.runs[0]
                            else:
                                target_run = paragraph.add_run()
                            
                            # Set the text
                            target_run.text = new_text
                            
                            # Set font to TH SarabunPSK size 16
                            target_run.font.name = 'TH SarabunPSK'
                            target_run.font.size = Pt(16)
                            
                            # Copy other formatting from template run if available
                            if template_run is not None:
                                target_run.font.bold = template_run.font.bold
                                target_run.font.italic = template_run.font.italic
                                target_run.font.underline = template_run.font.underline
                                if template_run.font.color.rgb:
                                    target_run.font.color.rgb = template_run.font.color.rgb
            
            # Helper function to find all tables recursively
            def replace_in_all_tables(element, replacements):
                """Recursively replace in all tables including nested ones"""
                if hasattr(element, 'tables'):
                    for table in element.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                # Replace in cell paragraphs
                                for paragraph in cell.paragraphs:
                                    replace_in_paragraph(paragraph, replacements)
                                
                                # Recursively process nested tables
                                replace_in_all_tables(cell, replacements)
            
            # Replace in main paragraphs
            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph, replacements)
            
            # Replace in all tables (including nested)
            replace_in_all_tables(doc, replacements)
            
            # Save filled document to temp folder
            filled_filename = f"filled_{report_name}"
            filled_path = os.path.join(self.temp_report_dir, filled_filename)
            
            # Remove old file if exists
            if os.path.exists(filled_path):
                os.remove(filled_path)
            
            doc.save(filled_path)
            
            # Get recorder (group_id) from logged-in user
            recorder_id = 0
            if self.main_controller and hasattr(self.main_controller, 'logged_in_user_info'):
                user_info = self.main_controller.logged_in_user_info
                if user_info and isinstance(user_info, dict):
                    recorder_id = user_info.get('group_id', 0)
            
            # Save to database
            success, report_form_id, message = self.report_form_service.create_report_form(
                lab_order_id=lab_order_id,
                location="",
                comment="",
                state=1,  # Changed from 0 to 1
                status=1,
                recorder=recorder_id,  # Use group_id from logged-in user
                approver=0,
                room_id=5  # Parasitology
            )
            
            if success:
                self.current_report_form_id = report_form_id
                # DEBUG: Uncomment for debugging
                # print(f"Report form created with ID: {report_form_id}")
            else:
                # DEBUG: Uncomment for debugging
                # print(f"Failed to create report form: {message}")
                pass
            
            return filled_path
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            # DEBUG: Uncomment for debugging
            # print(f"ERROR in fill_form_data:")
            # print(error_details)
            return None
    
    def on_save_button_clicked(self):
        """Event handler เมื่อกดปุ่มบันทึก - บันทึกไฟล์ .docx"""
        if not self.filled_word_path or not os.path.exists(self.filled_word_path):
            QMessageBox.warning(
                self.view,
                "แจ้งเตือน",
                "กรุณากดปุ่ม FORM PREVIEW ก่อนบันทึก"
            )
            return
        
        try:
            # Open file dialog to select save location
            default_filename = f"Parasite_Report_{self.selected_lab_order_id}.docx"
            
            file_path, _ = QFileDialog.getSaveFileName(
                self.view,
                "บันทึกไฟล์ฟอร์ม",
                default_filename,
                "Word Documents (*.docx)"
            )
            
            if file_path:
                # Ensure file has .docx extension
                if not file_path.lower().endswith('.docx'):
                    file_path += '.docx'
                
                # Copy filled word file to selected location
                shutil.copy2(self.filled_word_path, file_path)
                
                # Update location in database
                if self.current_report_form_id:
                    success, message = self.report_form_service.update_report_form(
                        report_id=self.current_report_form_id,
                        location=file_path
                    )
                    
                    if success:
                        # DEBUG: Uncomment for debugging
                        # print(f"Report form location updated: {file_path}")
                        pass
                    else:
                        # DEBUG: Uncomment for debugging
                        # print(f"Failed to update location: {message}")
                        pass
                
                QMessageBox.information(
                    self.view,
                    "สำเร็จ",
                    f"บันทึกไฟล์สำเร็จ\n{file_path}"
                )
        
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            # DEBUG: Uncomment for debugging
            # print(f"ERROR in on_save_button_clicked:")
            # print(error_details)
            QMessageBox.critical(
                self.view,
                "ข้อผิดพลาด",
                f"ไม่สามารถบันทึกไฟล์ได้:\n{str(e)}"
            )
    
    def load_barcode_data(self):
        """โหลดข้อมูลบาร์โค้ดจาก API"""
        self.is_loading = True
        try:
            result = self.send_lab_service.get_received_labs_to_day(str(self.room_id), self.current_offset, self.limit)
            if result and 'job_progress' in result:
                self.has_more_data = result.get('has_more', False)
                for item in result['job_progress']:
                    self._add_row_to_table(item)
                    self.all_data.append(item)
                self.current_offset += len(result['job_progress'])
            else:
                self.has_more_data = False
        except Exception as e:
            # DEBUG: Uncomment for debugging
            # print(f"Load Barcode Error: {e}")
            QMessageBox.warning(self.view, "ข้อผิดพลาด", f"ไม่สามารถโหลดข้อมูลบาร์โค้ดได้: {str(e)}")
        finally:
            self.is_loading = False
    
    def _add_row_to_table(self, item):
        """เพิ่มข้อมูลลงในตาราง"""
        lab_id = str(item.get('lab_order_id', '')).zfill(12)
        item_obj = QStandardItem(lab_id)
        item_obj.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        self.barcode_model.appendRow([item_obj])
    
    def on_scroll(self, value):
        """Event handler สำหรับ infinite scroll"""
        scrollbar = self.view.ui.barcode_para_tableView.verticalScrollBar()
        if value >= scrollbar.maximum() - 10:
            if self.has_more_data and not self.is_loading:
                if not self.view.ui.barcode_para_lineEdit.text():
                    self.load_more_data()
    
    def load_more_data(self):
        """โหลดข้อมูลเพิ่มเติม"""
        self.load_barcode_data()
    
    def on_barcode_double_clicked(self, index):
        """Event handler เมื่อ double click ที่ตาราง"""
        if index.isValid():
            row = index.row()
            barcode = self.barcode_model.item(row, 0).text()
            self.view.ui.barcode_para_lineEdit.setText(barcode)
            self.selected_lab_order_id = barcode
            # DEBUG: Uncomment for debugging
            # print(f"Selected Barcode: {barcode}")
    
    def on_select_barcode_clicked(self):
        """Event handler เมื่อกดปุ่มเลือกบาร์โค้ด"""
        barcode = self.view.ui.barcode_para_lineEdit.text().strip()
        if barcode:
            self.selected_lab_order_id = barcode
            # Show success (optional, can be removed for cleaner UX)
            # QMessageBox.information(self.view, "สำเร็จ", f"เลือกบาร์โค้ด: {barcode}")
            # DEBUG: Uncomment for debugging
            # print(f"Selected Barcode: {barcode}")
        else:
            QMessageBox.warning(self.view, "แจ้งเตือน", "กรุณากรอกหรือเลือกบาร์โค้ด")
    
    def on_barcode_search(self):
        """Event handler เมื่อกด Enter ในช่องค้นหา"""
        barcode = self.view.ui.barcode_para_lineEdit.text().strip()
        if barcode:
            self.search_by_barcode(barcode)
    
    def search_by_barcode(self, barcode):
        """ค้นหาบาร์โค้ด"""
        self.barcode_model.removeRows(0, self.barcode_model.rowCount())
        self.view.ui.barcode_para_tableView.scrollToTop()
        
        try:
            result = self.send_lab_service.get_received_labs_by_barcode(barcode, str(self.room_id))
            
            if result and result.get('found', False):
                for item in result['job_progress']:
                    self._add_row_to_table(item)
                # Show success (optional)
                # QMessageBox.information(self.view, "สำเร็จ", result.get('message', "พบข้อมูล"))
            else:
                QMessageBox.warning(self.view, "ไม่พบข้อมูล", "ไม่พบบาร์โค้ดนี้ในระบบ")
        except Exception as e:
            QMessageBox.warning(self.view, "ข้อผิดพลาด", f"เกิดข้อผิดพลาดในการค้นหา: {str(e)}")
    
    def reset_lazy_loading_state(self):
        """รีเซ็ตสถานะการโหลดข้อมูล"""
        self.current_offset = 0
        self.is_loading = False
        self.has_more_data = True
        self.all_data = []
    
    def reload_data(self):
        """โหลดข้อมูลใหม่"""
        self.barcode_model.removeRows(0, self.barcode_model.rowCount())
        self.reset_lazy_loading_state()
        self.view.ui.barcode_para_lineEdit.clear()
        self.load_barcode_data()
    
    def get_selected_color(self):
        """ดึงค่า Color ที่เลือก"""
        return self.view.ui.comboBox.currentText()
    
    def get_selected_consistency(self):
        """ดึงค่า Consistency ที่เลือก"""
        return self.view.ui.comboBox_2.currentText()
    
    def get_selected_method(self):
        """ดึงค่า Method ที่เลือก"""
        return self.view.ui.comboBox_3.currentText()
    
    def get_selected_barcode(self):
        """ดึงค่าบาร์โค้ดที่เลือก"""
        return self.selected_lab_order_id
    
    def clear_form(self):
        """ล้างค่าในฟอร์ม"""
        self.view.ui.comboBox.setCurrentIndex(-1)
        self.view.ui.comboBox_2.setCurrentIndex(-1)
        self.view.ui.comboBox_3.setCurrentIndex(-1)
        self.view.ui.barcode_para_lineEdit.clear()
        self.selected_lab_order_id = None
    
    def convert_form_to_pdf(self, word_path, report_name):
        """แปลงฟอร์ม Word เป็น PDF และแสดงผล"""
        try:
            # สร้างชื่อไฟล์ PDF
            pdf_filename = f"preview_{report_name.replace('.docx', '.pdf')}"
            self.pdf_output_path = os.path.join(self.temp_report_dir, pdf_filename)
            
            # ลบไฟล์เก่าถ้ามี
            if os.path.exists(self.pdf_output_path):
                os.remove(self.pdf_output_path)
            
            # เริ่มการแปลงใน Thread
            self.convert_thread = ConvertPDFThread(word_path, self.pdf_output_path)
            self.convert_thread.finished_signal.connect(self.on_conversion_finished)
            self.convert_thread.error_signal.connect(self.on_conversion_error)
            self.convert_thread.start()
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            # DEBUG: Uncomment for debugging
            # print(f"ERROR in convert_form_to_pdf:")
            # print(error_details)
            QMessageBox.critical(
                self.view, 
                "ข้อผิดพลาด", 
                f"ไม่สามารถเริ่มแปลงไฟล์ได้:\n{str(e)}"
            )
    
    def on_conversion_finished(self):
        """ทำงานเมื่อแปลงไฟล์สำเร็จ"""
        if self.progress_dialog:
            self.progress_dialog.setValue(100)
            self.progress_dialog.setLabelText("เสร็จสิ้น!")
            self.progress_dialog.close()
        self.display_pdf_preview()
        self.convert_thread = None
    
    def on_conversion_error(self, error_msg):
        """ทำงานเมื่อแปลงไฟล์ Error"""
        QMessageBox.critical(
            self.view, 
            "ข้อผิดพลาด", 
            f"ไม่สามารถแปลงไฟล์ได้:\n{error_msg}"
        )
        self.convert_thread = None
    
    def display_pdf_preview(self):
        """แสดงตัวอย่าง PDF ใน WebEngineView"""
        if self.pdf_output_path and os.path.exists(self.pdf_output_path):
            try:
                abs_pdf_path = os.path.abspath(self.pdf_output_path)
                
                settings = self.view.ui.preview_word_webEngineView.page().settings()
                settings.setAttribute(QWebEngineSettings.WebAttribute.PluginsEnabled, True)
                settings.setAttribute(QWebEngineSettings.WebAttribute.PdfViewerEnabled, True)
                
                with open(abs_pdf_path, "rb") as f:
                    pdf_data = f.read()
                    base64_pdf = base64.b64encode(pdf_data).decode('utf-8')
                
                pdf_filename = os.path.basename(abs_pdf_path)
                html_content = self.create_pdf_viewer_html(base64_pdf, pdf_filename)
                self.view.ui.preview_word_webEngineView.setHtml(html_content, QUrl("http://localhost"))
                
            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                print(f"ERROR in display_pdf_preview:")
                print(error_details)
                QMessageBox.warning(
                    self.view, 
                    "ข้อผิดพลาด", 
                    f"ไม่สามารถแสดงตัวอย่าง PDF ได้:\n{str(e)}"
                )
        else:
            QMessageBox.warning(
                self.view,
                "ข้อผิดพลาด",
                f"ไม่พบไฟล์ PDF ที่แปลงแล้ว\n{self.pdf_output_path}"
            )
    
    def create_pdf_viewer_html(self, base64_data, filename="document.pdf"):
        """สร้าง HTML สำหรับแสดง PDF"""
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{filename}</title>
            <style>
                body, html {{ margin: 0; padding: 0; height: 100%; overflow: hidden; background-color: #525659; }}
                embed {{ width: 100%; height: 100%; border: none; }}
            </style>
        </head>
        <body>
            <embed 
                src="data:application/pdf;base64,{base64_data}#toolbar=1&navpanes=0&scrollbar=1&page=1&view=FitH" 
                type="application/pdf" width="100%" height="100%"
            >
        </body>
        </html>
        """
