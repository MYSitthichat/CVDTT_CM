from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
import mariadb
import re
import os
import tempfile
from database import get_db_connection

router = APIRouter(tags=["Receive_lab_order"])

# DEBUG MODE - เปลี่ยนเป็น True เพื่อแสดงข้อมูลทั้งหมด, False เพื่อแสดงแค่ 3 ฟิลด์
DEBUG = True  # เปลี่ยนค่านี้เป็น True/False

# Pydantic models สำหรับรับข้อมูล
class ReceiveLabRequest(BaseModel):
    lab_order_id: int
    receive_from_room: int
    comment_for_sample: str = ""
    sample_status: str = "ปกติ"
    updater_id: int  # employee_id ของคนที่ login


class RejectLabRequest(BaseModel):
    lab_order_id: int
    receive_from_room: int
    comment_for_sample: str = ""
    sample_status: str = "เสียหาย/ไม่ปกติ"
    updater_id: int  # employee_id ของคนที่ login


class ExportTemplateRequest(BaseModel):
    lab_order_id: int
    template_path: str
    template_name: str
    output_filename: str


@router.get("/get_lab_order/detail")
def get_lab_order_detail(lab_order_id: str, room_id: str = ""):
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    try:
        cursor = conn.cursor()
        if room_id:
            query = """
                SELECT 
                    lab_order.id AS lab_order_id,
                    sample_registration.dtime,
                    sample_registration.speed,
                    sample_registration.sample_inspection,
                    sample_registration.id AS sample_id,
                    lab_order.room_id
                FROM lab_order
                INNER JOIN sample_registration
                    ON lab_order.sample_id = sample_registration.id
                WHERE lab_order.id = ? AND 
                lab_order.room_id = ? AND lab_order.status = 1
                """
            cursor.execute(query, (lab_order_id, room_id))
    except mariadb.Error as e:
        print(f"Query Error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve lab order detail")
    finally:
        conn.close()




@router.get("/get_lab_order/to_day")
def get_lab_order_to_day(room_id: str, offset: int = 0, limit: int = 50):
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    try:
        cursor = conn.cursor()
        cursor.execute("""
                SELECT 
                    sample_registration.dtime,
                    lab_order.id AS lab_order_id,
                    sample_registration.speed,
                    sample_registration.sample_inspection
                FROM lab_order
                INNER JOIN sample_registration 
                    ON lab_order.sample_id = sample_registration.id
                WHERE lab_order.room_id = ? AND lab_order.status = 1
                ORDER BY lab_order.id DESC 
                LIMIT ? OFFSET ?
                """, (room_id, limit, offset))
        groups = [{"time": row[0], "lab_order_id": row[1], "speed": row[2], "sample_inspection": row[3]} for row in cursor]
        
        cursor.execute("SELECT COUNT(*) FROM lab_order WHERE room_id = ? AND status = 1", (room_id,))
        total_count = cursor.fetchone()[0]
        
        return {
            "job_progress": groups,
            "total": total_count,
            "offset": offset,
            "limit": limit,
            "has_more": (offset + limit) < total_count
        }
    except mariadb.Error as e:
        print(f"Query Error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve job progress")
    finally:
        conn.close()

@router.get("/get_lab_order/barcode")
def get_lab_order_by_barcode(barcode: str, room_id: str = ""):
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    try:
        cursor = conn.cursor()
        
        # Convert barcode to lab_order_id (remove leading zeros)
        try:
            lab_order_id = int(barcode)
        except ValueError:
            return {
                "job_progress": [],
                "total": 0,
                "found": False,
                "message": "Barcode ไม่ถูกต้อง"
            }
        
        # Build query with room_id filter if provided
        if room_id:
            query = """
                SELECT 
                    sample_registration.dtime,
                    lab_order.id AS lab_order_id,
                    sample_registration.speed,
                    sample_registration.sample_inspection,
                    sample_registration.id AS sample_id,
                    lab_order.room_id
                FROM lab_order
                INNER JOIN sample_registration 
                    ON lab_order.sample_id = sample_registration.id
                WHERE lab_order.id = ? 
                    AND lab_order.room_id = ? 
                    AND lab_order.status = 1
                ORDER BY lab_order.id DESC 
                """
            cursor.execute(query, (lab_order_id, room_id))
        else:
            # Admin mode - no room restriction
            query = """
                SELECT 
                    sample_registration.dtime,
                    lab_order.id AS lab_order_id,
                    sample_registration.speed,
                    sample_registration.sample_inspection,
                    sample_registration.id AS sample_id,
                    lab_order.room_id
                FROM lab_order
                INNER JOIN sample_registration 
                    ON lab_order.sample_id = sample_registration.id
                WHERE lab_order.id = ? AND lab_order.status = 1
                ORDER BY lab_order.id DESC 
                """
            cursor.execute(query, (lab_order_id,))
        
        results = cursor.fetchall()
        
        if not results:
            if room_id:
                return {
                    "job_progress": [],
                    "total": 0,
                    "found": False,
                    "message": "ไม่พบข้อมูล Barcode นี้ในห้องของคุณ หรือไม่มีสิทธิ์เข้าถึง"
                }
            else:
                return {
                    "job_progress": [],
                    "total": 0,
                    "found": False,
                    "message": "ไม่พบข้อมูล Barcode นี้"
                }
        
        groups = [{
            "time": row[0], 
            "lab_order_id": row[1], 
            "speed": row[2], 
            "sample_inspection": row[3],
            "sample_id": row[4],
            "room_id": row[5]
        } for row in results]
        
        return {
            "job_progress": groups,
            "total": len(groups),
            "found": True,
            "message": f"พบข้อมูล {len(groups)} รายการ"
        }
    except mariadb.Error as e:
        print(f"Query Error: {e}")
        raise HTTPException(status_code=500, detail="Failed to search by barcode")
    finally:
        conn.close()

@router.get("/get_lab_order/details")
def get_lab_order_details(lab_order_id: str, room_id: str):
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    
    cursor = None
    try:
        cursor = conn.cursor()
        
        # แปลง lab_order_id เป็น int
        try:
            order_id = int(lab_order_id)
        except ValueError:
            return {
                "success": False,
                "message": "Lab Order ID ไม่ถูกต้อง"
            }
        
        # ดึงข้อมูล Lab Order และ Sample Registration
        sql_order = """
            SELECT 
                lo.dtime, 
                lo.id, 
                sr.species, 
                ri.code, 
                ri.nickname, 
                sr.keep_method, 
                sr.speed,
                sr.name,
                sr.opd_number,
                sr.sex,
                sr.age_year,
                sr.age_month,
                sr.age_day,
                sr.breed,
                sr.sample_type,
                cr.project_name,
                sr.id AS sample_id,
                lo.room_id,
                sr.sample_inspection
            FROM lab_order lo
            LEFT JOIN sample_registration sr ON lo.sample_id = sr.id 
            LEFT JOIN room_information ri ON lo.room_id = ri.id 
            LEFT JOIN case_registration cr ON sr.case_id = cr.id
            WHERE lo.id = %s AND lo.status = 1
        """
        
        if room_id:
            sql_order += " AND lo.room_id = %s"
            cursor.execute(sql_order, (order_id, room_id))
        else:
            cursor.execute(sql_order, (order_id,))
        
        order_result = cursor.fetchone()
        
        if not order_result:
            return {
                "success": False,
                "message": "ไม่พบข้อมูล Lab Order นี้"
            }
        
        # แปลง tuple เป็น dict
        order_data = {
            "dtime": order_result[0],
            "lab_order_id": order_result[1],
            "species": order_result[2],
            "room_code": order_result[3],
            "room_nickname": order_result[4],
            "keep_method": order_result[5],
            "speed": order_result[6],
            "name": order_result[7],
            "opd_number": order_result[8],
            "sex": order_result[9],
            "age_year": order_result[10],
            "age_month": order_result[11],
            "age_day": order_result[12],
            "breed": order_result[13],
            "sample_type": order_result[14],
            "project_name": order_result[15],
            "sample_id": order_result[16],
            "room_id": order_result[17],
            "sample_inspection": order_result[18]
        }
        
        # ดึงรายการตรวจตาม room_id
        test_items = []
        sample_id = order_data["sample_id"]
        room_id_val = order_data["room_id"]
        room_code = order_data["room_code"] if order_data["room_code"] else ""
        if room_id_val == 5:  # Parasite
            sql_test = """SELECT * FROM lab_parasite_biology WHERE sample_id = %s"""
            cursor.execute(sql_test, (sample_id,))
            test_data = cursor.fetchone()
            if test_data:
                raw_test_data = test_data[3:]  # Get all columns after id, sample_id, dtime
                for i in range(0, len(raw_test_data), 3):
                    test_name = raw_test_data[i] if i < len(raw_test_data) else ""
                    test_state = int(raw_test_data[i+1]) if i+1 < len(raw_test_data) and raw_test_data[i+1] is not None else 0
                    test_price = int(raw_test_data[i+2]) if i+2 < len(raw_test_data) and raw_test_data[i+2] is not None else 0
                    # เฉพาะรายการที่ state > 0 (state คือจำนวนเลย)
                    if test_state > 0 and test_name:
                        # ลบตัวเลขในวงเล็บออก เช่น "PCV (50)" -> "PCV"
                        # แปลงเป็น string ก่อนใช้ re.sub เพื่อป้องกัน TypeError
                        test_name_str = str(test_name) if test_name is not None else ""
                        clean_name = re.sub(r'\s*\(\d+\)\s*$', '', test_name_str).strip()
                        test_items.append({
                            "test_name": clean_name,
                            "test_amount": test_state  # ใช้ state เป็นจำนวน
                        })
        
        elif room_id_val == 2:  # Bacteria
            sql_test = """SELECT * FROM lab_bacteria_biology WHERE sample_id = %s"""
            cursor.execute(sql_test, (sample_id,))
            
            col_names = [desc[0] for desc in cursor.description]
            test_data = cursor.fetchone()
            
            if test_data:
                
                try:
                    # Process preparation_p1-p21 (use state as amount)
                    for i in range(1, 22):  # p1 to p21
                        name_idx = col_names.index(f'preparation_p{i}_name')
                        state_idx = col_names.index(f'preparation_p{i}_amount')
                        
                        test_name = test_data[name_idx] if test_data[name_idx] else ""
                        test_state = int(test_data[state_idx]) if test_data[state_idx] is not None else 0
                        
                        if test_state > 0 and test_name:
                            test_name_str = str(test_name) if test_name is not None else ""
                            clean_name = re.sub(r'\s*\(\d+\)\s*$', '', test_name_str).strip()
                            test_items.append({
                                "test_name": clean_name,
                                "test_amount": test_state
                            })
                    
                    # Process drug_sensitivity1-41 (use state as amount)
                    for i in range(1, 42):  # 1 to 41
                        name_idx = col_names.index(f'drug_sensitivity{i}_name')
                        state_idx = col_names.index(f'drug_sensitivity{i}_state')
                        
                        test_name = test_data[name_idx] if test_data[name_idx] else ""
                        test_state = int(test_data[state_idx]) if test_data[state_idx] is not None else 0
                        
                        if test_state > 0 and test_name:
                            test_name_str = str(test_name) if test_name is not None else ""
                            clean_name = re.sub(r'\s*\(\d+\)\s*$', '', test_name_str).strip()
                            test_items.append({
                                "test_name": clean_name,
                                "test_amount": test_state
                            })
                    
                    # Process bacteria_id1-12 (use state as amount)
                    for i in range(1, 13):  # 1 to 12
                        name_idx = col_names.index(f'bacteria_id{i}_name')
                        state_idx = col_names.index(f'bacteria_id{i}_state')
                        
                        test_name = test_data[name_idx] if test_data[name_idx] else ""
                        test_state = int(test_data[state_idx]) if test_data[state_idx] is not None else 0
                        
                        if test_state > 0 and test_name:
                            test_name_str = str(test_name) if test_name is not None else ""
                            clean_name = re.sub(r'\s*\(\d+\)\s*$', '', test_name_str).strip()
                            test_items.append({
                                "test_name": clean_name,
                                "test_amount": test_state
                            })
                    
                    # Process lab_request1-5 (use state as amount)
                    for i in range(1, 6):  # 1 to 5
                        name_idx = col_names.index(f'lab_request{i}_name')
                        state_idx = col_names.index(f'lab_request{i}_state')
                        
                        test_name = test_data[name_idx] if test_data[name_idx] else ""
                        test_state = int(test_data[state_idx]) if test_data[state_idx] is not None else 0
                        
                        if test_state > 0 and test_name:
                            test_name_str = str(test_name) if test_name is not None else ""
                            clean_name = re.sub(r'\s*\(\d+\)\s*$', '', test_name_str).strip()
                            test_items.append({
                                "test_name": clean_name,
                                "test_amount": test_state
                            })
                except Exception as e:
                    # If error occurs during parsing, return what we have so far
                    print(f"Error parsing bacteria data: {e}")
        
        elif room_id_val == 8:  # Molecular Biology
            sql_test = """SELECT * FROM lab_molecular_biology WHERE sample_id = %s"""
            cursor.execute(sql_test, (sample_id,))
            test_data = cursor.fetchone()
            if test_data:
                # Parse molecular test data (columns 3 onwards, every 3 columns = test_name, test_amount, test_price)
                # Structure: test1_name, test1_amount, test1_price, test2_name, test2_amount, test2_price, ...
                # สำหรับ Molecular: ไม่มี state, amount คือจำนวน (0=ไม่เลือก, >0=เลือกและเป็นจำนวน)
                raw_test_data = test_data[3:]
                for i in range(0, len(raw_test_data), 3):
                    test_name = raw_test_data[i] if i < len(raw_test_data) else ""
                    test_amount = int(raw_test_data[i+1]) if i+1 < len(raw_test_data) and raw_test_data[i+1] is not None else 0
                    test_price = int(raw_test_data[i+2]) if i+2 < len(raw_test_data) and raw_test_data[i+2] is not None else 0
                    # เฉพาะรายการที่ amount > 0 (amount คือจำนวนเลย)
                    if test_amount > 0 and test_name:
                        # ลบตัวเลขในวงเล็บออก - แปลงเป็น string ก่อนเพื่อป้องกัน error
                        test_name_str = str(test_name) if test_name is not None else ""
                        clean_name = re.sub(r'\s*\(\d+\)\s*$', '', test_name_str).strip()
                        test_items.append({
                            "test_name": clean_name,
                            "test_amount": test_amount  # ใช้ amount เป็นจำนวน
                        })
        
        return {
            "success": True,
            "order_data": order_data,
            "test_items": test_items
        }
        
    except mariadb.Error as e:
        print(f"Query Error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve lab order details")
    
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


@router.post("/receive_lab_order")
def receive_lab_order(request: ReceiveLabRequest):
    """
    บันทึกข้อมูลการรับแลป
    - room_action_status = 1 (รับแลป)
    - sample_status: "ปกติ" หรือ "เสียหาย/ไม่ปกติ"
    - ดึง case_id จาก lab_order โดยอัตโนมัติ
    """
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    
    cursor = None
    try:
        cursor = conn.cursor()
        
        # ดึงข้อมูล lab_order และ case_id
        sql_check = """
            SELECT lo.id, sr.case_id 
            FROM lab_order lo
            LEFT JOIN sample_registration sr ON lo.sample_id = sr.id
            WHERE lo.id = %s AND lo.status = 1
        """
        cursor.execute(sql_check, (request.lab_order_id,))
        result = cursor.fetchone()
        
        if not result:
            return {
                "success": False,
                "message": "ไม่พบ Lab Order นี้"
            }
        
        case_id = result[1]  # ดึง case_id จาก sample_registration
        
        # บันทึกข้อมูลการรับแลป
        sql = """
            INSERT INTO lab_receive_detail 
            (lab_order_id, case_id, receive_from_room, comment_for_sample, room_action_status,  updater) 
            VALUES (%s, %s, %s, %s, 1, %s)
        """
        
        cursor.execute(sql, (request.lab_order_id, case_id, request.receive_from_room, request.comment_for_sample, request.updater_id))
        conn.commit()
        
        return {
            "success": True,
            "message": "บันทึกการรับแลปสำเร็จ",
            "lab_order_id": request.lab_order_id
        }
        
    except mariadb.Error as e:
        if conn:
            conn.rollback()
        print(f"Database Error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to save receive lab data: {str(e)}")
    
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


@router.post("/reject_lab_order")
def reject_lab_order(request: RejectLabRequest):
    """
    ปฏิเสธการรับแลป
    - room_action_status = 0 (ปฏิเสธ)
    - ดึง case_id จาก lab_order โดยอัตโนมัติ
    """
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    
    cursor = None
    try:
        cursor = conn.cursor()
        
        # ดึงข้อมูล lab_order และ case_id
        sql_check = """
            SELECT lo.id, sr.case_id 
            FROM lab_order lo
            LEFT JOIN sample_registration sr ON lo.sample_id = sr.id
            WHERE lo.id = %s AND lo.status = 1
        """
        cursor.execute(sql_check, (request.lab_order_id,))
        result = cursor.fetchone()
        
        if not result:
            return {
                "success": False,
                "message": "ไม่พบ Lab Order นี้"
            }
        
        case_id = result[1]  # ดึง case_id จาก sample_registration
        
        # บันทึกข้อมูลการปฏิเสธแลป
        sql = """
            INSERT INTO lab_receive_detail 
            (lab_order_id, case_id, receive_from_room, comment_for_sample, room_action_status, updater) 
            VALUES (%s, %s, %s, %s, 0, %s)
        """
        
        cursor.execute(sql, (request.lab_order_id, case_id, request.receive_from_room, request.comment_for_sample, request.updater_id))
        conn.commit()
        
        return {
            "success": True,
            "message": "บันทึกการปฏิเสธแลปสำเร็จ",
            "lab_order_id": request.lab_order_id
        }
        
    except mariadb.Error as e:
        if conn:
            conn.rollback()
        print(f"Database Error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to save reject lab data: {str(e)}")
    
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


@router.get("/get_report_templates")
def get_report_templates(room_id: int = None):
    """
    ดึงรายการ report templates จาก database
    - หากมี room_id จะกรองเฉพาะ template ของห้องนั้น
    - หากไม่มี room_id จะดึงทั้งหมด
    """
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    
    cursor = None
    try:
        cursor = conn.cursor()
        
        if room_id:
            query = """
                SELECT id, report_name, room_id, report_path, updater 
                FROM report_information 
                WHERE room_id = ?
                ORDER BY report_name
            """
            cursor.execute(query, (room_id,))
        else:
            query = """
                SELECT id, report_name, room_id, report_path, updater 
                FROM report_information
                ORDER BY room_id, report_name
            """
            cursor.execute(query)
        
        results = cursor.fetchall()
        
        templates = []
        for row in results:
            template_info = {
                'id': row[0],
                'report_name': row[1],
                'room_id': row[2],
                'report_path': row[3],
                'updater': row[4]
            }
            templates.append(template_info)
        
        return {
            "success": True,
            "templates": templates,
            "count": len(templates)
        }
        
    except mariadb.Error as e:
        print(f"Query Error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve report templates: {str(e)}")
    
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()




@router.get("/get_template_data")
def get_template_data(lab_order_id: int):
    """
    ดึงข้อมูลจาก database สำหรับเติมลงใน template
    - ข้อมูล lab_receive_detail (report_id)
    - ข้อมูล sample_registration และ lab_order
    - ข้อมูล owner และ sender จาก customer
    """
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    
    cursor = None
    try:
        from datetime import datetime
        cursor = conn.cursor()
        
        # 1. ดึง id จาก lab_receive_detail (เลขที่รายงาน)
        query1 = """
            SELECT id, dtime
            FROM lab_receive_detail 
            WHERE lab_order_id = ? 
            ORDER BY id DESC 
            LIMIT 1
        """
        cursor.execute(query1, (lab_order_id,))
        result1 = cursor.fetchone()
        report_id = result1[0] if result1 else None
        receive_dtime = result1[1] if result1 else None
        
        # 2. ดึงข้อมูลจาก lab_order และ sample_registration
        query2 = """
            SELECT sr.collect_date, sr.id as case_id, sr.name, sr.species, sr.breed, 
                   sr.sex, sr.age_year, sr.age_month, sr.age_day, sr.sample_type, 
                   lo.id as lab_order_id, sr.case_id
            FROM lab_order lo
            LEFT JOIN sample_registration sr ON lo.sample_id = sr.id
            WHERE lo.id = ?
        """
        cursor.execute(query2, (lab_order_id,))
        result2 = cursor.fetchone()
        
        if result2:
            collect_date = result2[0]
            case_id = result2[1]
            animal_name = result2[2]
            species = result2[3]
            breed = result2[4]
            sex = result2[5]
            age_year = result2[6]
            age_month = result2[7]
            age_day = result2[8]
            sample_type = result2[9]
            registration_case_id = result2[11]
        else:
            raise HTTPException(status_code=404, detail="Lab order not found")
        
        # 3. ดึงข้อมูล owner และ sender
        owner_name = owner_phone = owner_email = owner_address = ''
        sender_name = sender_phone = sender_email = sender_address = ''
        
        if registration_case_id:
            query3 = """
                SELECT cr.owner_id, cr.sender_id,
                       o.name as owner_name, o.surname as owner_surname, 
                       o.phone as owner_phone, o.email as owner_email, 
                       o.contact_address as owner_address,
                       s.name as sender_name, s.surname as sender_surname,
                       s.phone as sender_phone, s.email as sender_email,
                       s.contact_address as sender_address
                FROM case_registration cr
                LEFT JOIN customer o ON cr.owner_id = o.id
                LEFT JOIN customer s ON cr.sender_id = s.id
                WHERE cr.id = ?
            """
            cursor.execute(query3, (registration_case_id,))
            result3 = cursor.fetchone()
            
            if result3:
                owner_name = f"{result3[2] or ''} {result3[3] or ''}".strip()
                owner_phone = result3[4] or ''
                owner_email = result3[5] or ''
                owner_address = result3[6] or ''
                
                sender_name = f"{result3[7] or ''} {result3[8] or ''}".strip()
                sender_phone = result3[9] or ''
                sender_email = result3[10] or ''
                sender_address = result3[11] or ''
        
        # 4. สร้างเลขที่ตัวอย่าง
        today = datetime.now()
        day_str = str(today.day)
        sample_number = f"D{day_str}-{lab_order_id}"
        
        # 5. จัดรูปแบบอายุ
        age_text = ''
        age_parts = []
        if age_year:
            age_parts.append(f"{age_year} ปี")
        if age_month:
            age_parts.append(f"{age_month} เดือน")
        if age_day:
            age_parts.append(f"{age_day} วัน")
        age_text = ' '.join(age_parts)
        
        # 6. แปลงวันที่
        formatted_collect_date = ''
        if collect_date:
            try:
                if isinstance(collect_date, str):
                    date_part = collect_date.split(' ')[0]
                    date_obj = datetime.strptime(date_part, '%Y-%m-%d')
                    formatted_collect_date = date_obj.strftime('%d/%m/%Y')
                else:
                    formatted_collect_date = str(collect_date)
            except:
                formatted_collect_date = str(collect_date)
        
        # ตรวจสอบ DEBUG MODE
        if DEBUG:
            # DEBUG = TRUE: แสดงข้อมูลทั้งหมด
            return {
                "success": True,
                "data": {
                    'report_id': str(report_id) if report_id else '',
                    'sample_number': sample_number,
                    'collect_date': formatted_collect_date,
                    'lab_order_id': lab_order_id,
                    'case_id': case_id,
                    'animal_name': animal_name if animal_name else '',
                    'species': species if species else '',
                    'breed': breed if breed else '',
                    'sex': sex if sex else '',
                    'age': age_text,
                    'sample_type': sample_type if sample_type else '',
                    'receive_dtime': str(receive_dtime) if receive_dtime else '',
                    'owner_name': owner_name,
                    'owner_phone': owner_phone,
                    'owner_email': owner_email,
                    'owner_address': owner_address,
                    'sender_name': sender_name,
                    'sender_phone': sender_phone,
                    'sender_email': sender_email,
                    'sender_address': sender_address
                }
            }
        else:
            # DEBUG = FALSE: แสดงเฉพาะ 3 ฟิลด์
            return {
                "success": True,
                "data": {
                    'report_id': str(report_id) if report_id else '',
                    'sample_number': sample_number,
                    'collect_date': formatted_collect_date,
                    # ฟิลด์อื่นๆ จะเป็นค่าว่าง
                    'lab_order_id': lab_order_id,
                    'case_id': case_id,
                    'animal_name': '',
                    'species': '',
                    'breed': '',
                    'sex': '',
                    'age': '',
                    'sample_type': '',
                    'receive_dtime': '',
                    'owner_name': '',
                    'owner_phone': '',
                    'owner_email': '',
                    'owner_address': '',
                    'sender_name': '',
                    'sender_phone': '',
                    'sender_email': '',
                    'sender_address': ''
                }
            }
        
    except mariadb.Error as e:
        print(f"Query Error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve template data: {str(e)}")
    
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


@router.post("/export_word_template")
def export_word_template(request: ExportTemplateRequest):
    """
    Export Word template โดยเติมข้อมูลจาก database
    - ดึงข้อมูลจาก get_template_data
    - เติมข้อมูลลงใน Word template
    - ส่งไฟล์กลับไปให้ client download
    """
    try:
        from docx import Document
        
        # 1. ดึงข้อมูลสำหรับเติมลงใน template
        template_data_result = get_template_data(request.lab_order_id)
        if not template_data_result.get('success'):
            raise HTTPException(status_code=404, detail="Cannot retrieve template data")
        
        data = template_data_result['data']
        
        # 2. สร้าง full path ของไฟล์ template
        source_file = os.path.join(request.template_path, request.template_name)
        
        if not os.path.exists(source_file):
            raise HTTPException(status_code=404, detail=f"Template file not found: {source_file}")
        
        # 3. เปิดไฟล์ template
        doc = Document(source_file)
        
        # 4. สร้าง mapping ของข้อมูล
        replacements = {
            'เลขที่รายงาน:': f"เลขที่รายงาน: {data.get('report_id', '')}",
            'วันที่รับตัวอย่าง:': f"วันที่รับตัวอย่าง: {data.get('collect_date', '')}",
            'เลขที่ตัวอย่าง:': f"เลขที่ตัวอย่าง: {data.get('sample_number', '')}",
            'พันธุ์:': f"พันธุ์: {data.get('breed', '')}",
            'เพศ:': f"เพศ: {data.get('sex', '')}",
        }
        
        # 5. แทนที่ข้อความใน paragraphs
        def replace_text_in_paragraph(paragraph):
            full_text = paragraph.text
            for key, value in replacements.items():
                if key in full_text:
                    original_font = None
                    if len(paragraph.runs) > 0:
                        original_font = paragraph.runs[0].font
                    
                    inline = paragraph.runs
                    for i in range(len(inline)-1, -1, -1):
                        p = inline[i]._element
                        p.getparent().remove(p)
                    
                    new_text = full_text.replace(key, value)
                    new_run = paragraph.add_run(new_text)
                    
                    if original_font:
                        new_run.font.name = original_font.name
                        new_run.font.size = original_font.size
                        new_run.font.bold = original_font.bold
                        new_run.font.italic = original_font.italic
                        new_run.font.underline = original_font.underline
                        new_run.font.color.rgb = original_font.color.rgb
                    break
        
        # แทนที่ใน paragraphs
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph)
        
        # 6. แทนที่ข้อความใน tables
        processed_cells = set()
        
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_text = ' '.join([cell.text for cell in row.cells])
                
                for cell_idx, cell in enumerate(row.cells):
                    cell_position = (row_idx, cell_idx)
                    cell_id = id(cell._element)
                    is_merged = False
                    for prev_cell_idx in range(cell_idx):
                        if id(row.cells[prev_cell_idx]._element) == cell_id:
                            is_merged = True
                            break
                    
                    if is_merged or cell_position in processed_cells:
                        continue
                    
                    processed_cells.add(cell_position)
                    
                    for para in cell.paragraphs:
                        cell_text = para.text.strip()
                        
                        # จัดการฟิลด์ที่อยู่ - หาแถวที่มี "ที่อยู่:" และตรวจสอบว่าเป็นแถวของเจ้าของหรือผู้ส่ง
                        if 'ที่อยู่:' in cell_text:
                            # ตรวจสอบว่าเป็นแถวก่อนหน้า (เจ้าของหรือผู้ส่ง)
                            is_owner_row = False
                            is_sender_row = False
                            
                            # ตรวจสอบแถวก่อนหน้า 1-2 แถว
                            if row_idx > 0:
                                prev_row_text = ' '.join([c.text for c in table.rows[row_idx - 1].cells])
                                if 'ชื่อเจ้าของ:' in prev_row_text:
                                    is_owner_row = True
                                elif 'ชื่อผู้ส่ง:' in prev_row_text:
                                    is_sender_row = True
                            
                            if is_owner_row:
                                # ที่อยู่ของเจ้าของ
                                for para_item in cell.paragraphs:
                                    for run in para_item.runs:
                                        run.text = ''
                                    if len(para_item.runs) > 0:
                                        para_item.runs[0].text = f"ที่อยู่: {data.get('owner_address', '')}"
                                    else:
                                        para_item.add_run(f"ที่อยู่: {data.get('owner_address', '')}")
                                    break
                                continue
                            elif is_sender_row:
                                # ที่อยู่ของผู้ส่ง
                                for para_item in cell.paragraphs:
                                    for run in para_item.runs:
                                        run.text = ''
                                    if len(para_item.runs) > 0:
                                        para_item.runs[0].text = f"ที่อยู่: {data.get('sender_address', '')}"
                                    else:
                                        para_item.add_run(f"ที่อยู่: {data.get('sender_address', '')}")
                                    break
                                continue
                        
                        # จัดการฟิลด์เจ้าของ
                        if 'ชื่อเจ้าของ:' in row_text or 'ชื่อ:' in row_text:
                            if cell_text.startswith('ชื่อเจ้าของ:'):
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ''
                                    if len(para.runs) > 0:
                                        para.runs[0].text = f"ชื่อเจ้าของ: {data.get('owner_name', '')}"
                                    else:
                                        para.add_run(f"ชื่อเจ้าของ: {data.get('owner_name', '')}")
                                    break
                            elif cell_text.startswith('โทร.'):
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ''
                                    if len(para.runs) > 0:
                                        para.runs[0].text = f"โทร. {data.get('owner_phone', '')}"
                                    else:
                                        para.add_run(f"โทร. {data.get('owner_phone', '')}")
                                    break
                            elif cell_text.startswith('E-mail:'):
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ''
                                    if len(para.runs) > 0:
                                        para.runs[0].text = f"E-mail: {data.get('owner_email', '')}"
                                    else:
                                        para.add_run(f"E-mail: {data.get('owner_email', '')}")
                                    break
                        
                        # จัดการฟิลด์ผู้ส่ง
                        elif 'ชื่อผู้ส่ง:' in row_text:
                            if cell_text.startswith('ชื่อผู้ส่ง:'):
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ''
                                    if len(para.runs) > 0:
                                        para.runs[0].text = f"ชื่อผู้ส่ง: {data.get('sender_name', '')}"
                                    else:
                                        para.add_run(f"ชื่อผู้ส่ง: {data.get('sender_name', '')}")
                                    break
                            elif cell_text.startswith('โทร.'):
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ''
                                    if len(para.runs) > 0:
                                        para.runs[0].text = f"โทร. {data.get('sender_phone', '')}"
                                    else:
                                        para.add_run(f"โทร. {data.get('sender_phone', '')}")
                                    break
                            elif cell_text.startswith('E-mail:'):
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ''
                                    if len(para.runs) > 0:
                                        para.runs[0].text = f"E-mail: {data.get('sender_email', '')}"
                                    else:
                                        para.add_run(f"E-mail: {data.get('sender_email', '')}")
                                    break
                        
                        # จัดการฟิลด์สัตว์
                        if cell_text.startswith('ชนิดสัตว์:'):
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.text = ''
                                if len(para.runs) > 0:
                                    para.runs[0].text = f"ชนิดสัตว์: {data.get('species', '')}"
                                else:
                                    para.add_run(f"ชนิดสัตว์: {data.get('species', '')}")
                                break
                        elif cell_text.startswith('ชื่อสัตว์:'):
                            # หา cell ถัดไปในแถวเดียวกัน
                            try:
                                # หา index ของ cell ปัจจุบัน
                                for idx, c in enumerate(row.cells):
                                    if c._element == cell._element:
                                        current_idx = idx
                                        break
                                
                                # เติมข้อมูลใน cell ถัดไป
                                if current_idx + 1 < len(row.cells):
                                    next_cell = row.cells[current_idx + 1]
                                    for next_para in next_cell.paragraphs:
                                        # ล้างข้อมูลเดิม
                                        for run in next_para.runs:
                                            run.text = ''
                                        # เติมข้อมูลใหม่
                                        if len(next_para.runs) > 0:
                                            next_para.runs[0].text = data.get('animal_name', '')
                                        else:
                                            next_para.add_run(data.get('animal_name', ''))
                                        break
                            except Exception as e:
                                print(f"Error filling animal name: {e}")
                                pass
                        elif cell_text.startswith('ชนิดตัวอย่าง:'):
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.text = ''
                                if len(para.runs) > 0:
                                    para.runs[0].text = f"ชนิดตัวอย่าง: {data.get('sample_type', '')}"
                                else:
                                    para.add_run(f"ชนิดตัวอย่าง: {data.get('sample_type', '')}")
                                break
                        elif cell_text.startswith('อายุ:'):
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.text = ''
                                if len(para.runs) > 0:
                                    para.runs[0].text = f"อายุ: {data.get('age', '')}"
                                else:
                                    para.add_run(f"อายุ: {data.get('age', '')}")
                                break
                        
                        # แทนที่ข้อความปกติ
                        replace_text_in_paragraph(para)
        
        # 7. บันทึกไฟล์ลง temporary location
        with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as tmp_file:
            output_path = tmp_file.name
            doc.save(output_path)
        
        # 8. ส่งไฟล์กลับไปให้ client
        return FileResponse(
            path=output_path,
            filename=request.output_filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            background=None  # Don't delete file immediately, let OS handle it
        )
        
    except Exception as e:
        print(f"Export Error: {e}")
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to export template: {str(e)}")